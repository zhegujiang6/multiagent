"""生成多智能体协作项目分析报告 Word 文档."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from pathlib import Path


def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2F5496')

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, 'D6E4F0')

    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.4
    style.paragraph_format.space_after = Pt(6)
    # Set CJK font
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ── Title ──
    title = doc.add_heading('全链路客户服务与工单闭环协同智能体', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('—— 多智能体协作项目技术分析报告 ——')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'生成日期：{datetime.date.today().isoformat()}').font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. 项目概述
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('一、项目概述', level=1)

    doc.add_paragraph(
        '本项目是一个基于 Multi-Agent + LLM 的智能客服运营解决方案。系统由 6 个专业 AI Agent '
        '协同工作，覆盖从客户接入、意图识别、情绪分析、知识检索、工单创建到闭环解决的全链路。'
        '项目采用 LangGraph 进行 Agent 状态编排，通过 RAG（检索增强生成）技术从知识库检索答案，'
        '并结合 Websocket 实现实时通信，为客户提供智能化的自助服务体验。'
    )

    doc.add_heading('核心功能', level=2)
    features = [
        ('智能对话', '基于 LLM 的自然语言理解与生成，支持多轮对话'),
        ('意图识别', '13 分类意图识别 + 实体提取，精准理解用户诉求'),
        ('情绪感知', '5 分类情绪分析 + 触发词检测，实时感知用户情绪状态'),
        ('知识检索', '基于 Qdrant 向量数据库的 RAG（检索增强生成），50+ FAQ + 10 SOP'),
        ('工单管理', '完整的工单生命周期状态机（8 状态 + 12 条合法转换）'),
        ('SLA 追踪', '智能 SLA 计算与自动升级，30 秒轮询检查'),
        ('自动转人工', '情绪愤怒/绝望、法律威胁、用户明确要求时自动升级'),
        ('客户画像', 'VIP 识别、服务策略定制、个性化服务等级'),
        ('自进化闭环', '知识缺口检测 → 人工填充 → 自动提取 → 自动审核 → 自动发布'),
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}：')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 2. 技术架构
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('二、技术架构', level=1)

    doc.add_heading('2.1 技术栈总览', level=2)

    tech_rows = [
        ['LLM 大模型', 'OpenAI 兼容 API（支持 Claude / Qwen 等）', '异步客户端，15s 超时，自动重试'],
        ['Agent 编排', 'LangGraph 0.2+ (StateGraph)', '有状态工作流，条件路由，并行执行'],
        ['后端框架', 'Python 3.13 + FastAPI', 'REST API + WebSocket 双通道'],
        ['向量数据库', 'Qdrant v1.12', '4 个 Collection，Cosine 相似度检索'],
        ['Embedding', 'sentence-transformers (本地) 或 OpenAI API', 'paraphrase-multilingual-MiniLM-L12-v2'],
        ['关系数据库', 'PostgreSQL 16 + pgvector', '7 张核心表，SQLAlchemy ORM + Alembic 迁移'],
        ['缓存/消息', 'Redis 7', '会话缓存 + SLA Pub/Sub + WebSocket 状态'],
        ['前端框架', 'React 18 + TypeScript + Vite', 'TailwindCSS 4 + Zustand 状态管理'],
        ['部署方式', 'Docker Compose', '5 个容器：postgres + redis + qdrant + backend + frontend'],
    ]
    add_styled_table(doc, ['技术层', '选型', '说明'], tech_rows)

    doc.add_heading('2.2 基础设施架构', level=2)
    doc.add_paragraph(
        '系统由 5 个 Docker 容器组成：\n'
        '• PostgreSQL 16 (pgvector/pgvector:pg16) — 存储会话、消息、工单、知识库元数据\n'
        '• Redis 7 (redis:7-alpine) — 缓存、发布订阅、WebSocket 状态管理\n'
        '• Qdrant (qdrant/qdrant:v1.12.0) — 向量存储与语义检索\n'
        '• Backend (Python/FastAPI) — Agent 编排、API 服务、WebSocket\n'
        '• Frontend (React/Nginx) — 客户聊天窗口 + 坐席工作台'
    )

    doc.add_heading('2.3 数据库模型', level=2)
    model_rows = [
        ['conversations', '会话记录', 'customer_id, channel, status, sentiment_trend, meta_info'],
        ['messages', '消息记录', 'conversation_id, role, content, content_type, meta_info'],
        ['users', '用户/客户', 'external_id, name, email, tier, tags'],
        ['tickets', '工单', 'display_id, title, category, priority, status, sla_deadline, assigned_to'],
        ['ticket_events', '工单事件', 'ticket_id, from_status, to_status, triggered_by, comment'],
        ['knowledge_articles', '知识库', 'title, content, category, status, source_ticket_id, effectiveness_score'],
        ['agent_runs', 'Agent 执行日志', 'conversation_id, agent_name, input/output_summary, latency_ms, error'],
    ]
    add_styled_table(doc, ['表名', '用途', '核心字段'], model_rows)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 3. Agent 体系
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('三、Agent 体系详解', level=1)

    doc.add_heading('3.1 Agent 总览', level=2)
    doc.add_paragraph(
        '系统共有 6 个专业 Agent，各自承担不同的职责。所有 Agent 均继承自 BaseAgent 抽象基类，'
        '通过统一的 execute(state) → dict 接口进行交互，共享 AgentState（TypedDict）作为数据契约。'
    )

    agent_rows = [
        ['Orchestrator\n（编排器）', '核心调度中心', '始终运行',
         '场景理解、任务分解、子 Agent 编排、结果聚合\n基于 LangGraph StateGraph 实现'],
        ['Intent Classifier\n（意图分类器）', '意图识别与实体提取', '每轮对话',
         '13 分类意图识别（faq/refund/tech_support/complaint 等）\n提取订单号、商品名、金额等实体'],
        ['Sentiment Analyzer\n（情绪分析器）', '情绪感知与预警', '每轮对话',
         '5 分类情绪分析（satisfied/neutral/dissatisfied/angry/desperate）\n触发词检测（投诉/315/律师等）\n历史情绪趋势追踪'],
        ['FAQ Agent\n（FAQ 问答）', '知识检索与答案生成', 'FAQ/产品信息类问题',
         '多路径 RAG 检索（Query Rewriting + 多 Collection + 结果融合）\n基于检索到的知识生成结构化答案\n知识缺口检测 + 使用追踪'],
        ['Ticket Agent\n（工单创建）', '工单信息提取与创建', '退款/技术/物流类问题',
         '自动提取工单信息（标题/分类/优先级/描述/建议部门）\n综合考虑情绪和 VIP 等级升级优先级\n自动计算 SLA 响应和解决截止时间'],
        ['Profile Enricher\n（画像增强）', '客户画像查询', '已识别用户',
         '从数据库查询客户信息（会员等级/历史/标签）\n提供差异化服务策略（VIP 热情主动 / 普通友好专业）\n匿名用户返回默认画像'],
    ]
    add_styled_table(doc, ['Agent', '职责', '触发时机', '核心能力'], agent_rows)

    doc.add_heading('3.2 Agent 实现细节', level=2)

    # --- Orchestrator ---
    doc.add_heading('3.2.1 Orchestrator（编排器）', level=3)
    doc.add_paragraph(
        'Orchestrator 是整个系统的"中枢神经系统"，基于 LangGraph StateGraph 实现。'
        '它定义一个包含 9 个节点的有状态工作流（workflow），通过条件边在节点之间路由。'
    )
    doc.add_paragraph(
        '工作流节点：\n'
        '① parallel_preprocess — 并行执行意图分类 + 情绪分析（asyncio.gather）\n'
        '② enrich_profile — 客户画像增强（查询数据库）\n'
        '③ route — 基于 LLM 的路由决策（4 条路径：FAQ/工单/转人工/直接回复）\n'
        '④ faq_answer — FAQ Agent 回答\n'
        '⑤ create_ticket — 工单创建 Agent\n'
        '⑥ escalate_to_human — 转人工处理\n'
        '⑦ direct_response — 直接回复（问候/闲聊）\n'
        '⑧ synthesize_response — 回复合成（共情前缀 + VIP 后缀）'
    )
    doc.add_paragraph(
        '关键实现机制：\n'
        '• 并行预处理：Intent Classifier 和 Sentiment Analyzer 通过 asyncio.gather 并行执行，节省约 50% 时间\n'
        '• LLM 路由：使用快速模型（FAST_MODEL）进行路由决策，JSON 结构化输出\n'
        '• 回退路由：当 LLM 路由解析失败时，自动降级为基于规则的回退路由\n'
        '• Pipeline 超时：整体 pipeline 25s 超时，超时自动转人工\n'
        '• 流式处理优化：如果 streaming 层已预处理（意图+情绪+画像），Orchestrator 跳过重复计算'
    )

    # --- Intent Classifier ---
    doc.add_heading('3.2.2 Intent Classifier（意图分类器）', level=3)
    doc.add_paragraph(
        '使用 LLM 对用户的最后一条消息进行 13 分类意图识别，同时提取关键实体信息。\n\n'
        '支持的意图：faq / product_info / order_inquiry / order_modify / refund / '
        'tech_support / complaint / account / payment / shipping / membership / chitchat / other\n\n'
        '输出结构：\n'
        '• label: 意图标签\n'
        '• confidence: 置信度 0.0-1.0\n'
        '• entities: 提取的实体列表 [{name, value}]\n'
        '• sub_intents: 次要意图列表\n'
        '• is_ambiguous: 是否模糊\n\n'
        '异常处理：LLM 调用失败时，降级为 {"label": "faq", "confidence": 0.3}'
    )

    # --- Sentiment Analyzer ---
    doc.add_heading('3.2.3 Sentiment Analyzer（情绪分析器）', level=3)
    doc.add_paragraph(
        '分析用户消息的情绪状态，支持 5 类情绪标签和触发词检测。\n\n'
        '情绪标签：satisfied / neutral / dissatisfied / angry / desperate\n\n'
        '关键能力：\n'
        '• 触发词检测："投诉"、"曝光"、"律师"、"315"、"消协"、"再也不用了"、"叫你们经理"等\n'
        '• 情绪趋势评估：基于历史情绪状态判断 improving / stable / declining / first_message\n'
        '• 情绪升级联动：angry/desperate 或检测到触发词 → 自动触发转人工\n\n'
        '异常处理：LLM 调用失败时，降级为 neutral'
    )

    # --- FAQ Agent ---
    doc.add_heading('3.2.4 FAQ Agent（FAQ 问答 + 自进化）', level=3)
    doc.add_paragraph(
        'FAQ Agent 是最复杂的 Agent 之一，集成了 RAG 检索、答案生成、知识缺口检测和知识使用追踪。\n\n'
        '检索策略（多路径检索）：\n'
        '① Query Rewriting：先用 LLM 将用户口语化问题改写为精准检索查询\n'
        '② 多 Collection 搜索：并行搜索 FAQ + SOP + 历史工单解决方案（根据意图决定搜索范围）\n'
        '③ 双路检索融合：原始 query + 改写 query 分别检索，按 title 去重，按 score 排序\n'
        '④ 单次 8s 超时：Qdrant 检索超时不阻塞，返回空结果\n\n'
        '自进化机制：\n'
        '• 知识缺口检测：best_score < 0.5 → 创建 gap 记录（后台异步）→ 人工专家审核填充\n'
        '• 知识使用追踪：答案生成后 → 增量更新源文章的 usage_count 和 effectiveness_score\n'
        '• 有效性评分：滚动平均算法 (old_score × (n-1) + outcome) / n\n\n'
        '异常处理：LLM 答案生成失败 → 直接使用检索到的原始知识内容作为回答'
    )

    # --- Ticket Agent ---
    doc.add_heading('3.2.5 Ticket Agent（工单创建）', level=3)
    doc.add_paragraph(
        '从对话上下文中自动提取工单信息，计算优先级和 SLA 截止时间。\n\n'
        '工单分类：order_inquiry / refund / tech_support / complaint / account / shipping / other\n\n'
        '优先级规则：\n'
        '• P0（紧急）：系统故障、安全事件、VIP 重大投诉、法律威胁\n'
        '• P1（高）：支付失败、订单丢失、投诉升级、情绪绝望\n'
        '• P2（中）：退货问题、物流异常、账户问题\n'
        '• P3（低）：使用咨询、发票申请、一般 FAQ\n\n'
        '优先级升级规则：\n'
        '• 情绪 desperate + 触发词 → P1 升至 P0\n'
        '• 情绪 angry → P2/P3 升至 P1\n'
        '• VIP 用户 → 优先级自动提升一档\n\n'
        'SLA 计算：\n'
        '• P0: 响应 15 分钟 / 解决 1 小时\n'
        '• P1: 响应 30 分钟 / 解决 4 小时\n'
        '• P2: 响应 1 小时 / 解决 8 小时\n'
        '• P3: 响应 4 小时 / 解决 24 小时\n\n'
        '异常处理：LLM 提取失败 → 使用最后一条消息的前 30 字符作为标题，默认 P2 优先级'
    )

    # --- Profile Enricher ---
    doc.add_heading('3.2.6 Profile Enricher（画像增强）', level=3)
    doc.add_paragraph(
        '从 PostgreSQL 数据库查询客户信息，实现差异化服务。\n\n'
        '查询逻辑：通过 customer_id 或 external_id 匹配 User 表\n\n'
        '会员等级差异化：\n'
        '• VIP：优先处理，授权额度翻倍，语气热情主动，自动升级工单优先级\n'
        '• Premium：服务质量高于普通用户，关注满意度\n'
        '• Standard：按标准流程服务\n'
        '• Anonymous（匿名访客）：使用默认画像\n\n'
        '异常处理：数据库查询失败 → 降级为 standard 默认画像'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 4. Agent 调用流程
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('四、Agent 之间的调用流程', level=1)

    doc.add_heading('4.1 整体流程（同步模式 — REST API）', level=2)
    doc.add_paragraph(
        '用户通过 REST API 发送消息时的完整处理流程如下：\n\n'
        '┌─────────────────────────────────────────────────────────────┐\n'
        '│ 1. POST /api/v1/conversations/{id}/messages                 │\n'
        '│    ↓                                                        │\n'
        '│ 2. 保存用户消息到 PostgreSQL (role="customer")              │\n'
        '│    ↓                                                        │\n'
        '│ 3. 获取最近 10 条历史消息，构建 AgentState                  │\n'
        '│    ↓                                                        │\n'
        '│ 4. 初始化 AgentRegistry → 创建 5 个子 Agent 实例            │\n'
        '│    ↓                                                        │\n'
        '│ 5. 创建 OrchestratorAgent → 调用 orchestrator.run(state)    │\n'
        '│    │                                                        │\n'
        '│    ├─ [并行] IntentClassifier.execute(state)                │\n'
        '│    ├─ [并行] SentimentAnalyzer.execute(state)               │\n'
        '│    │   (asyncio.gather 并行执行，约 2-3s)                   │\n'
        '│    ↓                                                        │\n'
        '│    ├─ ProfileEnricher.execute(state)  (查询数据库，< 1s)   │\n'
        '│    ↓                                                        │\n'
        '│    ├─ LLM 路由决策 (FAST_MODEL, 约 1-2s)                   │\n'
        '│    ↓                                                        │\n'
        '│    ├─ 分支 1: faq_answer → FAQAgent.execute(state)          │\n'
        '│    │             (RAG 检索 + 答案生成, 约 3-5s)             │\n'
        '│    ├─ 分支 2: create_ticket → TicketAgent.execute(state)    │\n'
        '│    │             (工单信息提取, 约 1-2s)                     │\n'
        '│    ├─ 分支 3: escalate_to_human (直接返回转人工响应)       │\n'
        '│    ├─ 分支 4: direct_response (直接返回问候响应)           │\n'
        '│    ↓                                                        │\n'
        '│    └─ SynthesizeResponse (共情前缀 + VIP 后缀, < 0.5s)     │\n'
        '│    ↓                                                        │\n'
        '│ 6. 保存 Agent 回复到 PostgreSQL (role="agent")              │\n'
        '│    ↓                                                        │\n'
        '│ 7. 返回完整响应 (消息 + 工单信息 + 意图 + 情绪)             │\n'
        '└─────────────────────────────────────────────────────────────┘\n\n'
        '整体耗时：正常 8-12s，超时阈值 25s'
    )

    doc.add_heading('4.2 流式处理流程（WebSocket 模式）', level=2)
    doc.add_paragraph(
        'WebSocket 模式在同步模式的基础上增加了实时状态推送，让前端能展示每个 Agent 的执行进度：\n\n'
        'WebSocket 客户端 ←→ ws://host/api/v1/ws/chat?conversation_id={id}\n\n'
        '推送事件序列：\n'
        '① {"type": "message_saved"} → 用户消息已保存\n'
        '② {"type": "agent_status", "agent": "intent_classifier", "status": "started"}\n'
        '③ {"type": "agent_status", "agent": "sentiment_analyzer", "status": "started"}\n'
        '④ {"type": "agent_status", "agent": "intent_classifier", "status": "completed", "result": "refund"}\n'
        '⑤ {"type": "agent_status", "agent": "sentiment_analyzer", "status": "completed", "result": "neutral"}\n'
        '⑥ {"type": "agent_status", "agent": "profile_enricher", "status": "started"}\n'
        '⑦ {"type": "agent_status", "agent": "profile_enricher", "status": "completed"}\n'
        '⑧ {"type": "agent_status", "agent": "orchestrator", "status": "started"}\n'
        '⑨ {"type": "ticket_created", "draft": {...}} (如果是工单路径)\n'
        '⑩ {"type": "escalating", "eta_seconds": 30} (如果是转人工)\n'
        '⑪ {"type": "chat_message", "message": {...}} → 最终回复\n\n'
        'WebSocket 模式下，意图分析、情绪分析和画像增强在 streaming 层预先执行，'
        '然后把结果注入 AgentState 传给 Orchestrator，Orchestrator 检测到已有数据后跳过重复计算。'
    )

    doc.add_heading('4.3 Agent 注册与发现机制', level=2)
    doc.add_paragraph(
        'Agent 之间的调用通过 AgentRegistry（注册中心）实现，采用依赖注入模式：\n\n'
        '① 初始化阶段：AgentRegistry(llm_client).initialize() 创建所有 Agent 实例\n'
        '   - intent_classifier = IntentClassifierAgent(llm_client)\n'
        '   - sentiment_analyzer = SentimentAnalyzerAgent(llm_client)\n'
        '   - faq_agent = FAQAgent(llm_client)\n'
        '   - profile_enricher = ProfileEnricherAgent(llm_client)\n'
        '   - ticket_agent = TicketAgent(llm_client)\n\n'
        '② 调用阶段：registry.get("agent_name") 返回对应的 Agent 实例\n\n'
        '③ 数据传递：所有 Agent 通过统一的 AgentState (TypedDict) 共享数据\n'
        '   - 输入：state 包含 messages, conversation_id, customer_id\n'
        '   - 输出：每个 Agent 返回 dict[str, Any]，LangGraph 自动合并到 state\n'
        '   - 累积：agent_decisions 列表记录每个 Agent 的决策痕迹'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 5. 转人工机制
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('五、转人工机制（Human Handoff）', level=1)

    doc.add_paragraph(
        '系统设计了多层次、多触发条件的自动转人工机制，确保关键场景能及时获得人工介入。'
    )

    doc.add_heading('5.1 触发条件（4 类场景）', level=2)

    esc_rows = [
        ['情绪触发', '情绪 = angry 或 desperate\n或检测到高危触发词',
         '自动',
         '检测到愤怒/绝望情绪时立即升级，不等路由决策完成'],
        ['意图触发', '意图 = complaint + 法律威胁词\n（如：315、律师、消协、曝光）',
         '路由决策阶段',
         'OR 条件：complaint 意图 OR 法律威胁词'],
        ['用户明确要求', '用户说"转人工"/"叫你们经理"等',
         'LLM 路由决策',
         '自然语言理解，灵活匹配'],
        ['系统故障', 'Pipeline 超时 (25s)\nLLM 调用失败\n未知异常',
         '异常处理阶段',
         '超时/异常后自动降级为转人工，并附上错误信息'],
    ]
    add_styled_table(doc, ['触发类别', '具体条件', '检测阶段', '备注'], esc_rows)

    doc.add_heading('5.2 转人工流程', level=2)
    doc.add_paragraph(
        '当触发转人工时，系统执行以下步骤：\n\n'
        'Step 1：设置 should_escalate = True\n'
        'Step 2：生成转人工响应文本："您的咨询已升级为人工服务。我已经为您生成了问题摘要，'
        '专业客服人员正在赶来，预计等待时间不超过 30 秒..."\n'
        'Step 3：通过 WebSocket 推送 {"type": "escalating", "eta_seconds": 30} 事件\n'
        'Step 4：记录 agent_decisions 为 escalate_to_human\n'
        'Step 5：前端收到 escalating 事件后，自动跳转到排队等待界面\n'
        'Step 6：人工坐席在 Agent Workspace 中看到该会话，可以接管对话'
    )

    doc.add_heading('5.3 手动转人工', level=2)
    doc.add_paragraph(
        '除了自动转人工，系统也支持手动转人工：\n'
        '• REST API: POST /api/v1/conversations/{id}/escalate\n'
        '• 用户在前端聊天窗口点击"转人工"按钮\n'
        '• 返回 {status: "escalating", message: "正在为您转接人工客服，请稍候..."}'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 6. Agent 失败处理
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('六、Agent 执行失败处理机制', level=1)

    doc.add_paragraph(
        '系统为每个 Agent 和每个关键环节都设计了多层容错和降级策略，确保任何单一故障都不会导致'
        '用户得不到回复。'
    )

    doc.add_heading('6.1 分层失败处理策略', level=2)

    fail_rows = [
        ['LLM API 调用层', 'BaseAgent._call_llm()',
         '• asyncio.wait_for(timeout=15s)\n• 超时抛出 TimeoutError\n• 网络异常抛出原始异常',
         '• 上层 Agent 的 try/except 捕获\n• 使用默认/回退值'],
        ['Agent 执行层', '各 Agent.execute()',
         '• Intent Classifier 失败\n• Sentiment Analyzer 失败\n• FAQ Agent 失败\n• Ticket Agent 失败',
         '• 降级为 faq/chitchat (0.3 置信度)\n• 降级为 neutral\n'
         '• 返回知识缺口提示 + 建议转人工\n• 使用最后消息前 30 字符为标题，P2 优先级'],
        ['Orchestrator 层', 'orchestrator.run()',
         '• Pipeline 超时 (25s)\n• 未预期异常',
         '• 自动转人工："系统处理超时，正在为您转接人工客服..."\n• 同上'],
        ['服务层', 'process_message()\nprocess_message_streaming()',
         '• Orchestrator 整体异常\n• WebSocket 断开\n• RAG 检索超时',
         '• 自动转人工 + 记录 error 字段\n• WebSocket 发送 error 事件\n• 返回空结果，继续走 FAQ 回退逻辑'],
    ]
    add_styled_table(doc, ['层级', '关键组件', '可能失败', '降级策略'], fail_rows)

    doc.add_heading('6.2 LLM 调用层详解', level=2)
    doc.add_paragraph(
        'BaseAgent._call_llm() 是所有 Agent 调用 LLM 的统一入口，内置以下保护机制：\n\n'
        '1. 超时控制：每次 LLM 调用有 15 秒超时（asyncio.wait_for）\n'
        '2. 响应兼容：自动适配两种响应格式\n'
        '   - 标准 OpenAI ChatCompletion: resp.choices[0].message.content\n'
        '   - 阿里 MaaS 兼容模式: resp.text（choices 为 None）\n'
        '3. 日志记录：每次调用记录耗时和内容长度，便于排查问题\n'
        '4. 异常透传：超时/网络错误向上抛出，由调用方 Agent 处理'
    )

    doc.add_heading('6.3 路由决策层回退', level=2)
    doc.add_paragraph(
        '当 LLM 路由决策失败（JSON 解析异常）时，系统启动基于规则的回退路由 (_fallback_route)：\n\n'
        '回退规则（按优先级）：\n'
        '1. 情绪 angry/desperate 或存在触发词 → escalate_to_human\n'
        '2. 意图 complaint → escalate_to_human\n'
        '3. 意图 faq/product_info/account/chitchat/membership → faq_answer\n'
        '4. 意图 refund/order_inquiry/order_modify/tech_support/shipping/payment → create_ticket\n'
        '5. 默认 → faq_answer\n\n'
        '这套规则确保即使 LLM 完全不可用，系统仍能做出合理的路由决策。'
    )

    doc.add_heading('6.4 数据库/RAG 层容错', level=2)
    doc.add_paragraph(
        '• Profile Enricher：数据库查询失败 → 降级为 standard 默认画像\n'
        '• RAG 检索：Qdrant 连接超时/失败 → 返回空列表，FAQ Agent 自动进入知识缺口处理\n'
        '• Redis：缓存不可用 → 跳过缓存，直接查询（非致命）\n'
        '• PostgreSQL：连接失败 → FastAPI 抛出 500 错误，由全局异常处理器接管'
    )

    doc.add_heading('6.5 Pipeline 超时处理', level=2)
    doc.add_paragraph(
        '整个 Agent Pipeline 采用双层超时保护：\n\n'
        '内层（Orchestrator.run）：25 秒超时\n'
        '  - 正常流程 8-12s，25s 提供充裕的余量\n'
        '  - 超时 → 返回转人工响应："抱歉，系统处理您的请求时超时了..."\n\n'
        '外层（process_message_streaming）：30 秒超时\n'
        '  - 作为安全网，防止 Orchestrator 自身的超时机制失效\n'
        '  - 超时 → 同样返回转人工响应\n\n'
        '所有超时和异常场景，最终都会落到 should_escalate=True，保证用户永远能得到回复。'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 7. 工单状态机与 SLA
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('七、工单生命周期与 SLA 管理', level=1)

    doc.add_heading('7.1 工单状态机', level=2)
    doc.add_paragraph(
        '工单从创建到关闭经历严格的状态转换，每次状态变更都受状态机校验：\n\n'
        '状态流转图：\n'
        '  new → assigned → in_progress → resolved → closed\n'
        '                  ↓       ↑          ↑\n'
        '                pending / waiting     reopened\n\n'
        '共 8 个状态，12 条合法转换路径。非法转换（如 resolved → new）将被拒绝并抛出 ValueError。'
    )

    status_rows = [
        ['new', '新建', '新创建的工单，等待分配'],
        ['assigned', '已分配', '已分配给具体的客服人员'],
        ['in_progress', '处理中', '客服正在处理该问题'],
        ['pending', '挂起', '等待第三方/内部其他部门回复'],
        ['waiting', '等待中', '等待客户提供更多信息'],
        ['resolved', '已解决', '问题已解决，等待客户确认或自动关闭'],
        ['closed', '已关闭', '工单已关闭（终态，可重新打开）'],
        ['reopened', '重新打开', '已关闭的工单被客户重新激活'],
    ]
    add_styled_table(doc, ['状态', '中文名', '说明'], status_rows)

    doc.add_heading('7.2 SLA 自动追踪', level=2)
    doc.add_paragraph(
        'SLA（服务等级协议）追踪是系统后台持续运行的服务，每 30 秒轮询一次所有活跃工单：\n\n'
        '检查逻辑：\n'
        '• 50% SLA 时间消耗 → 发送预警（sla_warning_sent = True）\n'
        '• 80% SLA 时间消耗 → 自动升级优先级（P2→P1, P1→P0）+ 重新计算 SLA\n'
        '• 100% SLA 时间消耗 → 紧急升级（P0）+ Redis Pub/Sub 通知坐席\n\n'
        '通知机制：通过 Redis Pub/Sub 频道 "sla:escalations" 推送升级通知，'
        '前端坐席工作台实时接收并高亮显示。'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 8. 知识自进化闭环
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('八、知识自进化闭环', level=1)

    doc.add_paragraph(
        '系统实现了完整的知识自进化闭环（Self-Evolution），使知识库能够自动化地从实际对话中持续学习和成长。'
    )

    doc.add_heading('8.1 闭环流程', level=2)
    doc.add_paragraph(
        '┌──────────────────────────────────────────────────────────────┐\n'
        '│  Phase 1: 知识缺口检测（FAQ Agent）                         │\n'
        '│  检索最佳得分 < 0.5 → 后台创建 gap 记录 → 通知人工专家     │\n'
        '│  ↓                                                          │\n'
        '│  Phase 2: 人工专家填充（Knowledge Management UI）           │\n'
        '│  专家编写内容 → fill_gap() → 自动审批 + 发布到 Qdrant      │\n'
        '│  ↓                                                          │\n'
        '│  Phase 3: 工单解决自动提取（工单 resolved 时触发）          │\n'
        '│  LLM 提取 Q&A 对 → confidence ≥ 0.8 自动发布到 Qdrant      │\n'
        '│                   → confidence 0.6-0.8 存为草稿待审核       │\n'
        '│                   → confidence < 0.6 丢弃（低价值）          │\n'
        '│  ↓                                                          │\n'
        '│  Phase 4: 知识使用追踪（FAQ Agent 每次回答后）              │\n'
        '│  更新 usage_count → 计算 effectiveness_score（滚动平均）    │\n'
        '│  → 低效知识可以被标记为待审核或淘汰                        │\n'
        '└──────────────────────────────────────────────────────────────┘'
    )

    doc.add_heading('8.2 关键阈值', level=2)
    evo_rows = [
        ['AUTO_APPROVE_CONFIDENCE', '0.8', '自动审批并发布到 Qdrant，无需人工审核'],
        ['DRAFT_CONFIDENCE', '0.6', '置信度在此之上但不足 0.8 → 存为草稿，待人工审核'],
        ['MIN_ANSWER_SCORE', '0.5', 'RAG 检索得分低于此值 → 判定为知识缺口，触发 gap 检测'],
        ['DEFAULT_CONFIDENCE', '0.5', 'LLM 提取失败时的默认置信度'],
    ]
    add_styled_table(doc, ['阈值名称', '值', '作用'], evo_rows)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 9. API 体系
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('九、API 体系', level=1)

    doc.add_heading('9.1 REST API', level=2)
    api_rows = [
        ['POST', '/api/v1/conversations', '创建会话'],
        ['GET', '/api/v1/conversations/{id}', '获取会话详情'],
        ['GET', '/api/v1/conversations/{id}/messages', '获取消息历史'],
        ['POST', '/api/v1/conversations/{id}/messages', '发送消息（同步）'],
        ['POST', '/api/v1/conversations/{id}/escalate', '手动转人工'],
        ['POST', '/api/v1/conversations/{id}/close', '关闭会话并触发知识提取'],
        ['GET', '/api/v1/tickets', '工单列表（支持筛选）'],
        ['POST', '/api/v1/tickets', '创建工单'],
        ['GET', '/api/v1/tickets/{id}', '工单详情'],
        ['PATCH', '/api/v1/tickets/{id}', '更新工单状态'],
        ['POST', '/api/v1/tickets/{id}/comments', '添加工单评论'],
        ['GET', '/api/v1/admin/metrics', '控制台指标'],
        ['GET', '/api/v1/knowledge/*', '知识库 CRUD'],
        ['GET', '/api/v1/health', '健康检查'],
    ]
    add_styled_table(doc, ['方法', '路径', '说明'], api_rows)

    doc.add_heading('9.2 WebSocket', level=2)
    doc.add_paragraph(
        'WebSocket 端点：ws://host/api/v1/ws/chat?conversation_id={id}\n\n'
        '支持的消息类型：\n'
        '• ping/pong — 心跳保活\n'
        '• message — 发送用户消息，触发 Agent Pipeline\n'
        '• typing — 输入状态指示\n\n'
        '推送事件类型：\n'
        '• message_saved — 消息已持久化\n'
        '• agent_status — Agent 执行状态（started/completed）\n'
        '• ticket_created — 工单自动创建通知\n'
        '• escalating — 转接人工通知\n'
        '• chat_message — Agent 最终回复\n'
        '• error — 错误通知'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 10. 部署架构
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('十、部署架构与启动流程', level=1)

    doc.add_heading('10.1 Docker Compose 部署', level=2)
    doc.add_paragraph(
        'docker-compose.yml 定义了 5 个服务：\n\n'
        '① postgres (pgvector/pgvector:pg16)\n'
        '   - 端口 5432\n'
        '   - 健康检查：pg_isready 每 5s\n'
        '   - 初始化脚本：docker/postgres/init.sql\n\n'
        '② redis (redis:7-alpine)\n'
        '   - 端口 6379\n'
        '   - AOF 持久化 + Keyspace 通知\n\n'
        '③ qdrant (qdrant/qdrant:v1.12.0)\n'
        '   - 端口 6333 (HTTP) / 6334 (gRPC)\n\n'
        '④ backend (Python/FastAPI)\n'
        '   - 构建：backend/Dockerfile\n'
        '   - 启动命令：alembic upgrade head && uvicorn app.main:app --host 0.0.0.0 --port 8000\n'
        '   - 依赖条件：postgres healthy + redis healthy + qdrant started\n\n'
        '⑤ frontend (React/Nginx)\n'
        '   - 构建：frontend/Dockerfile\n'
        '   - 端口 80\n'
        '   - 依赖条件：backend started'
    )

    doc.add_heading('10.2 本地开发启动', level=2)
    doc.add_paragraph(
        '1. 启动基础设施：docker compose up -d postgres redis qdrant\n'
        '2. 初始化数据库：cd backend && alembic upgrade head\n'
        '3. 加载知识库：python -m app.rag.seed_data\n'
        '4. 启动后端：uvicorn app.main:app --reload --host 0.0.0.0 --port 8000\n'
        '5. 启动前端：cd frontend && npm install && npm run dev\n\n'
        '访问地址：\n'
        '• 客户聊天界面：http://localhost:5173/chat\n'
        '• 坐席工作台：http://localhost:5173/workspace\n'
        '• API 文档：http://localhost:8000/docs'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 11. 总结
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('十一、项目亮点总结', level=1)

    highlights = [
        ('多 Agent 协同', '6 个专业 Agent 各司其职，通过 LangGraph StateGraph 统一编排，'
         '并行执行缩短响应时间，条件路由实现智能分流。'),
        ('有状态编排', '基于 LangGraph 的 AgentState（TypedDict）作为所有 Agent 之间的共享数据契约，'
         'agent_decisions 提供完整的决策审计追踪。'),
        ('鲁棒的容错设计', '每个层面都有降级策略：LLM 超时 15s、Pipeline 超时 25s/30s、'
         'LLM 失败用回退规则、DB 失败用默认值、RAG 失败走知识缺口。'
         '任何单一故障都不会导致用户得不到回复。'),
        ('实时反馈', 'WebSocket 推送每个 Agent 的执行状态，用户和坐席都能看到处理进度，'
         '不再是"黑盒"等待。'),
        ('多层转人工', '4 类触发条件覆盖情绪、意图、用户要求、系统故障，确保关键场景不遗漏。'),
        ('知识自进化', '知识缺口自动检测 + 工单解决自动提取 + 置信度分级自动审核 + '
         '使用效果滚动追踪，形成完整闭环。'),
        ('SLA 自动管理', '后台 30s 轮询 + 3 级阈值（预警/升级/紧急）+ Redis Pub/Sub 实时通知，'
         '确保服务质量。'),
        ('灵活部署', 'Docker Compose 一键部署，5 个容器各司其职，健康检查确保启动顺序。'),
    ]
    for title, desc in highlights:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}：')
        run.bold = True
        p.add_run(desc)

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 文档结束 —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    return doc


if __name__ == '__main__':
    doc = build_document()
    output_dir = Path(__file__).resolve().parents[2] / 'docs' / 'design'
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / '多智能体协作项目-流程与技术分析报告.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')
