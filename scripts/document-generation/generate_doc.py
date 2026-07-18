from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from pathlib import Path

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── 样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_para(text, bold=False, size=None, color=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_code_block(text):
    """添加代码块样式段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_style(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para('生产环境故障分析 & 自动修复智能体', bold=True, size=26,
         color=(0x1A, 0x3C, 0x6E), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Production Fault Analysis & Auto-Remediation Agent', bold=False, size=14,
         color=(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

add_para('—— 基于 LLM + Multi-Agent 的智能运维方案 ——', bold=False, size=13,
         color=(0x33, 0x33, 0x33), alignment=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_para(f'文档版本：v1.0', size=10, color=(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f'日期：{datetime.date.today().isoformat()}', size=10, color=(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 目录页（手动）
# ═══════════════════════════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    ('1', '项目背景与目标', ''),
    ('2', '总体架构设计', ''),
    ('2.1', '架构全景图', ''),
    ('2.2', '核心设计原则', ''),
    ('3', '技术栈选型', ''),
    ('3.1', '技术栈总览', ''),
    ('3.2', '选型理由', ''),
    ('4', '核心模块设计', ''),
    ('4.1', '数据采集层（Observability Pipeline）', ''),
    ('4.2', '异常检测引擎（Anomaly Detection）', ''),
    ('4.3', '根因分析引擎（RCA — Root Cause Analysis）', ''),
    ('4.4', '自动修复执行器（Auto-Remediation Executor）', ''),
    ('4.5', '安全护栏 & 审批流（Guardrails & Approval）', ''),
    ('4.6', '知识库 & 长期记忆（Knowledge Base）', ''),
    ('4.7', 'LLM 推理中枢（LLM Orchestrator）', ''),
    ('5', 'Multi-Agent 协作机制', ''),
    ('6', '核心工作流程', ''),
    ('6.1', '故障全生命周期', ''),
    ('6.2', '关键路径时序', ''),
    ('7', '安全与合规设计', ''),
    ('8', '部署架构', ''),
    ('9', '迭代路线图', ''),
    ('10', '总结与展望', ''),
]
for num, title, _ in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if not num.startswith('    '):
        run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. 项目背景与目标
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. 项目背景与目标', level=1)

doc.add_heading('1.1 痛点分析', level=2)
pain_points = [
    ('MTTR 过长', '生产故障平均发现时间（MTTD）> 15min，平均修复时间（MTTR）> 60min，严重依赖 on-call 人员经验。'),
    ('告警风暴', '一次底层故障可能触发数百条告警，人工筛选根因困难，极易遗漏关键信号。'),
    ('经验孤岛', '资深 SRE 的排障知识存在于脑子里或散落的 runbook 中，人员离职 = 知识流失。'),
    ('重复劳动', '70% 以上的故障是已知问题，每次仍需要人工重复相同排查步骤。'),
    ('夜间疲劳', '凌晨 3 点的告警，响应速度和处理质量远低于白天。'),
]
for title, desc in pain_points:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_heading('1.2 项目目标', level=2)
goals = [
    '构建一个基于 LLM 的 Multi-Agent 系统，自动完成「故障发现 → 告警聚合 → 根因分析 → 修复执行 → 复盘报告」全链路。',
    '将 MTTR 从小时级压缩到分钟级，覆盖 80% 已知故障场景。',
    '提供安全护栏（Guardrails），确保自动修复操作可审计、可回滚、可干预。',
    '沉淀故障知识库，让每次排障都成为可复用的组织资产。',
]
for g in goals:
    p = doc.add_paragraph(f'• {g}')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. 总体架构设计
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. 总体架构设计', level=1)

doc.add_heading('2.1 架构全景图', level=2)

add_para('系统采用分层 + 微服务架构，核心分为 7 层，通过消息队列解耦，各层可独立扩缩容：')

# 架构图文字版
arch_text = """
┌─────────────────────────────────────────────────────────────────────┐
│                         🔔  接入层（Ingestion）                        │
│  Prometheus AlertManager  │  Grafana Alerting  │  ELK Watcher       │
│  自定义 Webhook  │  CloudWatch / Azure Monitor  │  Sentry / Datadog  │
└──────────────────────────────────┬──────────────────────────────────┘
                                   ▼
┌─────────────────────────────────────────────────────────────────────┐
│                     📊  数据采集 & 上下文增强层                         │
│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌────────────┐ │
│  │ 指标拉取  │ │ 日志采集  │ │ Trace获取 │ │ 拓扑发现  │ │ 变更事件   │ │
│  │ PromQL   │ │ Loki/ELK │ │ Jaeger   │ │ K8s/CMDB │ │ Git/Jenkins│ │
│  └──────────┘ └──────────┘ └──────────┘ └──────────┘ └────────────┘ │
└──────────────────────────────────┬──────────────────────────────────┘
                                   ▼
┌─────────────────────────────────────────────────────────────────────┐
│                     🧠  Multi-Agent 推理层（核心）                      │
│  ┌──────────────────────────────────────────────────────────────┐   │
│  │                  LLM Orchestrator (主控 Agent)                 │   │
│  │   Agentic Router → 任务分解 → 子Agent调度 → 结果聚合 → 决策     │   │
│  └──────────────────────────────────────────────────────────────┘   │
│       │              │               │              │               │
│       ▼              ▼               ▼              ▼               │
│  ┌─────────┐  ┌──────────┐  ┌──────────┐  ┌──────────────┐         │
│  │告警分析  │  │ 日志分析  │  │ 指标分析  │  │ 变更关联     │         │
│  │ Agent   │  │ Agent    │  │ Agent    │  │ Agent        │         │
│  └─────────┘  └──────────┘  └──────────┘  └──────────────┘         │
│       │              │               │              │               │
│       └──────────────┴───────────────┴──────────────┘               │
│                          │                                          │
│                          ▼                                          │
│               ┌──────────────────┐                                  │
│               │  根因分析 (RCA)   │ ← 多维度交叉验证                  │
│               │   综合研判 Agent  │                                  │
│               └──────────────────┘                                  │
└──────────────────────────────────┬──────────────────────────────────┘
                                   ▼
┌─────────────────────────────────────────────────────────────────────┐
│                     🛡️  安全护栏 & 人机协同层                          │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐               │
│  │ 操作风险分级  │  │  审批工作流   │  │  变更窗口检查  │               │
│  │ L1-L4       │  │  自动/人工    │  │  CR/冻结期    │               │
│  └──────────────┘  └──────────────┘  └──────────────┘               │
└──────────────────────────────────┬──────────────────────────────────┘
                                   ▼
┌─────────────────────────────────────────────────────────────────────┐
│                     🔧  自动修复执行层                                 │
│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌────────────┐ │
│  │ 服务重启  │ │ 流量切换  │ │ 扩容缩容  │ │ 配置回滚  │ │ Runbook   │ │
│  │ K8s/SSH │ │ Nginx/   │ │ HPA/Karp │ │ GitOps/  │ │ 脚本执行   │ │
│  │          │ │  DNS切换  │ │ enter    │ │ ArgoCD   │ │           │ │
│  └──────────┘ └──────────┘ └──────────┘ └──────────┘ └────────────┘ │
└──────────────────────────────────┬──────────────────────────────────┘
                                   ▼
┌─────────────────────────────────────────────────────────────────────┐
│                     📝  复盘 & 知识沉淀层                              │
│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌────────────────────────┐  │
│  │ 自动生成  │ │ 知识库更新│ │ Runbook  │ │ 向量数据库（长期记忆）   │  │
│  │ 复盘报告  │ │ 向量化存储│ │  自动生成 │ │ 相似故障检索            │  │
│  └──────────┘ └──────────┘ └──────────┘ └────────────────────────┘  │
└─────────────────────────────────────────────────────────────────────┘
"""
add_code_block(arch_text)

doc.add_heading('2.2 核心设计原则', level=2)
principles = [
    ('渐进式自动化', 'L1（建议）→ L2（半自动）→ L3（自动+审批）→ L4（全自动），根据操作风险等级逐级开放。'),
    ('可观测性驱动', '所有决策基于可观测数据（指标/日志/链路/事件），LLM 不凭空猜测，而是分析数据得出结论。'),
    ('安全第一', '任何写操作必须经过 Guardrails 校验，高风险操作强制走审批流。所有操作可审计、可回滚。'),
    ('人在回路（Human-in-the-Loop）', '系统做分析和建议，最终高风险决策由人类确认。系统从确认中持续学习。'),
    ('知识沉淀', '每次故障处理都是知识库的增量更新，系统越用越聪明。'),
]
for title, desc in principles:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. 技术栈选型
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. 技术栈选型', level=1)

doc.add_heading('3.1 技术栈总览', level=2)

tech_headers = ['层级', '技术选型', '说明']
tech_rows = [
    ['LLM 推理', 'Claude API (Opus 4.8 / Sonnet 5)\n+ 本地 Qwen3-235B (Fallback)', 'Opus 负责复杂 RCA 推理；\nSonnet 负责告警初筛和常规分析；\n本地模型处理敏感数据，无需出网。'],
    ['Agent 框架', 'Claude Agent SDK / LangGraph\n+ 自研 Agentic Router', 'LangGraph 构建有状态、可检查点的\nMulti-Agent 工作流；\nSDK 提供原生 tool-use 能力。'],
    ['消息队列', 'Kafka / Redpanda', '告警事件流、Agent 间消息传递，\n保证高吞吐和消息不丢失。'],
    ['指标监控', 'Prometheus + VictoriaMetrics\n+ Grafana', '时序指标存储与查询，\nVictoriaMetrics 提供长期存储。'],
    ['日志平台', 'Loki + Elasticsearch', 'Loki 存结构化日志，低成本；\nES 存全文检索日志。'],
    ['链路追踪', 'Jaeger + OpenTelemetry', '分布式链路追踪，\n用于服务依赖分析和延迟定位。'],
    ['拓扑/CMDB', 'Neo4j (服务依赖图)\n+ Kubernetes API + Consul', '实时服务拓扑发现和依赖关系查询。'],
    ['变更事件', 'GitLab/GitHub Webhook\n+ Jenkins/ArgoCD Events', '关联近期变更，快速定位\n"哪个部署引入了故障"。'],
    ['向量数据库', 'Milvus / Qdrant', '故障知识库的语义存储和相似检索，\n支持跨故障案例的类比推理。'],
    ['工作流引擎', 'Temporal / Argo Workflows', '编排长时间运行的修复工作流，\n支持重试、超时、补偿（回滚）。'],
    ['安全护栏', 'Open Policy Agent (OPA)\n+ 自定义 Python Guardrails', '声明式策略，"允许/拒绝"每次操作，\n如：禁止在生产窗口外执行重启。'],
    ['审批/通知', 'PagerDuty / 企业微信 / Slack\n+ 自定义审批服务', '高风险操作推送审批，\nIM 通道交互式确认。'],
    ['修复执行', 'Kubernetes API + Ansible AWX\n+ Rundeck', 'K8s 原生操作（重启/扩缩容/切换）；\n物理机/VM 通过 Ansible 执行。'],
    ['存储', 'PostgreSQL (元数据/审计)\n+ MinIO/S3 (报告/日志归档)', 'PG 存工单、审计、审批记录；\n对象存储存复盘报告和故障快照。'],
    ['部署', 'Kubernetes + Helm\n+ ArgoCD (GitOps)', '所有服务容器化，\nGitOps 管理配置和部署。'],
]
add_table_with_style(tech_headers, tech_rows)

doc.add_heading('3.2 为什么选 Claude + LangGraph', level=2)
reasons = [
    ('Claude API', '业界领先的推理能力，特别是 Opus 4.8 在复杂因果推理上表现卓越；支持 200K 上下文窗口，可一次性输入大量日志/指标；原生 tool-use 与 MCP 协议支持，方便集成各种工具。'),
    ('LangGraph', '专为 Agent 工作流设计的状态图框架，支持循环、分支、并行、人工审批节点；内置 checkpoint 机制，工作流可暂停/恢复/回放；与 LangSmith 集成，可追踪每一步 Agent 决策。'),
    ('混合模型策略', '敏感数据用本地部署的 Qwen 模型处理，不上传外网；复杂推理用 Claude Opus，日常分析用 Claude Sonnet，兼顾效果与成本。'),
]
for title, desc in reasons:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. 核心模块设计
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. 核心模块设计', level=1)

# ── 4.1 ──
doc.add_heading('4.1 数据采集层（Observability Pipeline）', level=2)
add_para('职责：在故障告警触发后，自动围绕告警对象拉取 360° 上下文数据，构建 Agent 的"感知场"。')

data_sources = [
    ('告警源适配器', '统一接入 Prometheus AlertManager、Grafana、ELK Watcher、CloudWatch、Datadog 等告警事件，标准化为 CloudEvents 格式。'),
    ('上下文增强器', '根据告警中的 service/pod/namespace 等标签，自动拉取：最近 30min 指标时序、关联错误日志、最近 5 次部署变更、服务依赖拓扑子图。'),
    ('变更关联器', '查询 Git/CI/CD 系统，找出告警时间窗口内的所有变更事件（代码提交、配置变更、镜像升级），按相关性排序。'),
]
for title, desc in data_sources:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 4.2 ──
doc.add_heading('4.2 异常检测引擎（Anomaly Detection）', level=2)
add_para('职责：在多维数据中自动发现异常模式，不依赖固定阈值，而是用统计 + ML + LLM 三道防线。')

detection = [
    ('第一道：统计规则', '同比/环比突变检测（P99 延迟翻倍、错误率突增）、季节性分解（Prophet / STL）、同环比 + IQR 异常检测。速度快，覆盖面广。'),
    ('第二道：ML 模型', 'Isolation Forest / LSTM-AutoEncoder 做无监督异常检测；Prophet + 置信区间检测趋势偏离。处理复杂模式，减少规则维护成本。'),
    ('第三道：LLM 模式识别', '将时序数据的文字描述 + 日志摘要输入 LLM，让模型从语义层面判断是否是异常、属于什么类型（如"OOM 型"、"雪崩型"、"慢查询型"）。'),
]
for title, desc in detection:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 4.3 ──
doc.add_heading('4.3 根因分析引擎（RCA Agent）', level=2)
add_para('职责：融合多维度证据链，推理故障根因。这是整个系统最核心的模块。')

add_para('分析范式：Hypothesis-Driven RCA（假设驱动的根因分析）', bold=True)

rca_steps = [
    ('Step 1 — 故障画像', '基于告警 + 异常 + 拓扑信息，生成故障画像：影响范围（blast radius）、故障类型（服务/资源/配置/网络）、严重等级。'),
    ('Step 2 — 假设生成', 'LLM 根据故障画像 + 历史相似案例（从向量数据库检索 top-k），生成 3-5 个根因假设，每个假设附置信度。'),
    ('Step 3 — 证据收集', '每个假设派发专门的子 Agent 去收集支持/反对证据。如"内存泄漏假设"→拉取 OOM 日志 + 内存曲线 + GC 日志；"配置错误假设"→diff 最近配置变更。'),
    ('Step 4 — 交叉验证', '综合研判 Agent 汇总所有假设的支撑证据，排除矛盾的、更新置信度，输出最可能的根因 + 证据链。'),
    ('Step 5 — 修复建议', '根据确认的根因，从知识库检索对应的修复方案（Runbook），生成具体操作步骤和预期影响。'),
]
for title, desc in rca_steps:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 4.4 ──
doc.add_heading('4.4 自动修复执行器（Auto-Remediation Executor）', level=2)
add_para('职责：将 RCA 输出的修复方案转换为安全的、可审计的执行操作。')

executor = [
    ('操作原子化', '将复杂修复动作拆分为原子步骤，每步有明确的 pre-check、执行、post-check、回滚定义。'),
    ('执行引擎', 'Kubernetes 操作（kubectl / client-go）、Ansible（物理机/VM）、Terraform（基础设施）、自定义脚本（沙箱执行）。'),
    ('回滚机制', '每个操作步骤完成后记录状态快照，失败时自动执行逆操作。通过 Temporal Workflow 保证最终一致性。'),
    ('效果验证', '修复执行后自动验证：检查告警是否恢复、指标是否回到正常区间、用户流量是否正常。验证不通过自动回滚。'),
]
for title, desc in executor:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 4.5 ──
doc.add_heading('4.5 安全护栏 & 审批流（Guardrails & Approval）', level=2)
add_para('职责：确保自动修复不会造成二次故障，是系统的"刹车"和"方向盘"。')

add_para('操作风险分级：', bold=True)
risk_headers = ['风险等级', '操作示例', '审批策略', '示例']
risk_rows = [
    ['L1 - 只读', '查询日志、拉取指标、检索知识库', '全自动，无需审批', 'kubectl get / describe\nPromQL 查询'],
    ['L2 - 低风险', '重启单个 Pod、清理磁盘缓存', '自动执行，事后通知', 'kubectl rollout restart\n清理 /tmp'],
    ['L3 - 中风险', '流量切换、水平扩容、配置热更新', '自动提案 → 人工一键审批\n（IM 交互 / Web UI）', '调整 Service 权重\nHPA min/max 修改'],
    ['L4 - 高风险', '数据库主从切换、全量回滚\n网络策略变更', '强制人工审批 → 双人复核\n→ 变更窗口检查', 'MySQL Failover\nNginx 全量 reload'],
]
add_table_with_style(risk_headers, risk_rows)

# ── 4.6 ──
doc.add_heading('4.6 知识库 & 长期记忆（Knowledge Base）', level=2)
add_para('职责：存储故障案例、Runbook、修复经验，让系统越用越聪明。')

kb = [
    ('向量存储 (Milvus)', '故障描述、日志特征、根因结论的 Embedding。支持语义检索："当前这个 OOM 跟 3 周前那次的模式是不是一样？"。'),
    ('图谱存储 (Neo4j)', '服务依赖关系、"故障→根因→修复→服务"的因果关系图。支持图遍历查询："这个服务挂了会影响哪些上游？"。'),
    ('Runbook 库 (Git)', '经过验证的修复手册，以 Markdown + 参数化模板存储。LLM 检索到匹配的 Runbook 后，填充参数自动生成执行计划。'),
    ('反馈闭环', '人工审批时可以修正 Agent 的判断或选择更合适的修复方案 → 反馈信号回写知识库 → 更新 Embedding 或 Rank 权重。'),
]
for title, desc in kb:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 4.7 ──
doc.add_heading('4.7 LLM 推理中枢（LLM Orchestrator）', level=2)
add_para('职责：整个系统的"大脑"，负责理解故障、调度子 Agent、综合推理、做出决策。')

orchestrator = [
    ('Agentic Router', '根据告警类型（CPU 异常 / 内存异常 / 错误率 / 网络 / 存储）路由到不同的分析管道，选择最合适的 Agent 组合。'),
    ('Prompt 工程', '每个 Agent 有精心设计的 System Prompt，包含：角色定义、可用工具、输出格式、思维链（CoT）、注意事项。使用动态 Few-Shot（从知识库检索相似案例作为示例）。'),
    ('上下文窗口管理', '200K 上下文窗口虽然大，但要精打细算。策略：分层摘要（原始数据→结构化摘要→关键发现→最终结论），按需加载（先摘要后详细）。'),
    ('模型路由', '简单任务（告警分类、格式化）→ Sonnet；复杂推理（RCA 根因分析、跨系统关联）→ Opus 4.8；敏感数据本地处理 → Qwen3。'),
    ('结构化输出', '所有 Agent 输出使用 JSON Schema 约束（tool_choice / structured outputs），确保下游可程序化消费，避免解析失败。'),
]
for title, desc in orchestrator:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. Multi-Agent 协作机制
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Multi-Agent 协作机制', level=1)

add_para('本系统的 Multi-Agent 架构参考了 Anthropic 的 "Orchestrator-Workers" 模式以及 AutoGen / CrewAI 的多 Agent 对话范式，但做了面向生产环境的定制：')

agent_headers = ['Agent 角色', '触发条件', '输入', '输出', '使用的模型']
agent_rows = [
    ['Orchestrator\n(主控 Agent)', '任意告警到达', '告警事件 + 上下文', '任务分解计划\n子 Agent 调度指令', 'Claude Opus 4.8'],
    ['Alert Classifier\n(告警分类)', '告警进入系统', '告警原始 JSON + 标签', '告警分类、严重度\n路由目标', 'Claude Sonnet 5'],
    ['Metrics Analyst\n(指标分析)', '告警涉及性能指标', 'PromQL 查询结果\n时序数据', '异常模式描述\n关键时间点', 'Claude Sonnet 5'],
    ['Log Analyst\n(日志分析)', '告警涉及应用日志', 'Loki/ES 日志', '错误模式、关键堆栈\n时间线重建', 'Claude Sonnet 5'],
    ['Change Correlator\n(变更关联)', '告警窗口内存在变更', 'Git log / CI 记录', '可疑变更列表\n相关度评分', 'Claude Sonnet 5'],
    ['Topology Explorer\n(拓扑分析)', '需要了解服务依赖', 'Neo4j / K8s API', '影响范围图\n上下游依赖', '无需 LLM\n(图查询)'],
    ['RCA Synthesizer\n(根因综合)', '各子 Agent 返回结论', '所有子 Agent 报告', '根因假设 + 置信度\n证据链', 'Claude Opus 4.8'],
    ['Remediation Planner\n(修复规划)', 'RCA 确认根因', '根因 + Runbook 库', '修复方案步骤\n风险评估', 'Claude Opus 4.8'],
    ['Post-mortem Writer\n(复盘撰写)', '故障关闭后', '全流程记录 + 时间线', 'Markdown 复盘报告\n改进建议', 'Claude Sonnet 5'],
]
add_table_with_style(agent_headers, agent_rows)

add_para('协作模式：', bold=True)
collab = [
    ('Orchestrator-Workers', '主控 Agent 负责全局推理和决策，子 Agent 各司其职提供专项分析。这种模式比"Agent 自由对话"更可控、更确定。'),
    ('并行 + 串行混合', '告警分类、指标分析、日志分析、变更关联可以并行执行；RCA 综合必须在所有分析结果到位后串行执行；修复规划依赖 RCA 结果。'),
    ('共享上下文（Blackboard）', '所有 Agent 的中间结果写入共享的"黑板"（Redis + PG）。Orchestrator 从黑板读取所有子结论，进行综合判断。'),
    ('辩论机制（Debate）', '当 RCA Synthesizer 无法确定唯一根因时，可以启动 Debate 模式：让两个 Agent 分别持不同假设辩论，Orchestrator 作为裁判裁决。'),
]
for title, desc in collab:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. 核心工作流程
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. 核心工作流程', level=1)

doc.add_heading('6.1 故障全生命周期', level=2)

workflow_text = """
┌──────────┐    ┌──────────┐    ┌──────────┐    ┌──────────┐    ┌──────────┐
│ 1. 告警   │───▶│ 2. 上下文 │───▶│ 3. 异常   │───▶│ 4. 根因   │───▶│ 5. 修复   │
│    接入   │    │    增强   │    │    检测   │    │    分析   │    │    执行   │
└──────────┘    └──────────┘    └──────────┘    └──────────┘    └──────────┘
                                                                      │
                                                                      ▼
┌──────────┐    ┌──────────┐    ┌──────────┐    ┌──────────┐    ┌──────────┐
│10. 知识库 │◀───│ 9. 反馈   │◀───│ 8. 复盘   │◀───│ 7. 效果   │◀───│ 6. 审批   │
│    更新   │    │    学习   │    │    报告   │    │    验证   │    │    决策   │
└──────────┘    └──────────┘    └──────────┘    └──────────┘    └──────────┘
"""
add_code_block(workflow_text)

doc.add_heading('6.2 关键路径时序（典型场景：服务 OOM 故障）', level=2)

flow_headers = ['时间', '阶段', '执行者', '动作']
flow_rows = [
    ['T+0s', '告警接入', 'AlertManager → Kafka', 'Prometheus 检测到 payment-svc Pod 反复 OOMKilled，发出 Critical 告警。'],
    ['T+5s', '上下文增强', '数据采集层', '拉取 payment-svc 最近 30min 的：内存曲线、GC 日志、JVM 堆 dump 信息、最近 2 次部署 diff、上下游服务拓扑。'],
    ['T+10s', '异常检测', 'Metrics Analyst Agent', '分析内存曲线，发现内存呈单调递增，典型的"锯齿形"OOM 模式，判断为内存泄漏而非瞬时 spike。'],
    ['T+15s', '假设生成', 'Orchestrator', '检索向量库中与"内存泄漏+OOM+Java"相似的历史案例 Top-5，生成 3 个假设：\nH1: JVM 堆配置不足 (置信度 0.3)\nH2: 新版本代码引入内存泄漏 (置信度 0.6)\nH3: 上游流量激增导致 (置信度 0.1)'],
    ['T+25s', '证据收集', '多个子 Agent 并行', 'Log Analyst: 在最近一次部署后日志中出现了新的 WARN 日志"Connection pool not released"\nChange Correlator: 2h 前有一次 payment-svc 的代码发布\nMetrics Analyst: 上游流量无明显变化，排除 H3。'],
    ['T+30s', '根因确认', 'RCA Synthesizer', '综合证据：新版本中 Connection Pool 未正确释放 → H2 置信度提升至 0.92。输出结论："根因为 v2.4.1 代码变更引入的连接池泄漏，导致堆内存持续增长直至 OOM"。'],
    ['T+35s', '修复方案', 'Remediation Planner', '方案：回滚 payment-svc 到 v2.4.0 → 重启所有 Pod → 验证内存趋势恢复正常。操作等级 L3，需要人工审批。'],
    ['T+40s', '审批', 'SRE 值班人员', 'Slack 收到审批卡片，可视化展示根因证据和修复方案。SRE 一键 Approve。'],
    ['T+45s', '执行修复', 'Auto-Remediation', 'ArgoCD 触发回滚 → 等待 Pod 就绪 → 验证健康检查通过 → 观察 5min 内存曲线平稳。'],
    ['T+300s', '效果验证', '效果验证模块', '5min 观察窗口内：内存稳定在 60%，无 OOMKilled 事件，QPS 恢复正常。标记本次故障为已解决。'],
    ['T+310s', '复盘生成', 'Post-mortem Writer', '自动生成复盘报告：故障时间线、根因分析过程、修复操作记录、预防建议（修复 Connection Pool Bug、增加内存泄漏监控告警）。写入知识库。'],
]
add_table_with_style(flow_headers, flow_rows)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. 安全与合规设计
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. 安全与合规设计', level=1)

security = [
    ('操作沙箱', '所有 LLM 生成的命令/脚本必须在沙箱环境中预执行验证，不允许直接在生产环境执行未经校验的操作。使用 eBPF / seccomp 限制系统调用。'),
    ('变更窗口检查', '集成企业变更日历，高风险操作自动检查是否在维护窗口/变更冻结期。非窗口期的高风险操作拒绝执行。'),
    ('双人复核', 'L4 级别操作（数据库切换、全网流量切换）必须两名 SRE 同时审批。审批记录写入不可篡改的审计日志。'),
    ('敏感数据保护', '配置本地模型处理含敏感信息的日志数据；所有传输加密（mTLS）；向量数据库中存储的是脱敏后的 Embedding。'),
    ('审计跟踪', '每一个 Agent 的决策、每一次操作执行、每一次审批都记录完整的审计日志（PostgreSQL + 对象存储），支持按故障 ID 全链路回溯。'),
    ('权限最小化', '每个修复执行器仅有完成其任务所需的最小权限（K8s RBAC、AWS IAM Role）。Agent 之间使用独立的 Service Account。'),
    ('熔断机制', '连续 N 次修复执行失败 → 自动熔断，停止自动操作 → 升级为人工处理。防止自动修复造成雪崩。'),
]
for title, desc in security:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. 部署架构
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. 部署架构', level=1)

deploy_text = """
                          ┌──────────────────────────┐
                          │     Load Balancer         │
                          │   (Nginx / ALB / Traefik) │
                          └─────────────┬────────────┘
                                        │
                    ┌───────────────────┼───────────────────┐
                    ▼                   ▼                   ▼
          ┌─────────────┐     ┌─────────────┐     ┌─────────────┐
          │  K8s Node 1 │     │  K8s Node 2 │     │  K8s Node 3 │
          │             │     │             │     │             │
          │ ┌─────────┐ │     │ ┌─────────┐ │     │ ┌─────────┐ │
          │ │Ingestion│ │     │ │  Agent  │ │     │ │  Agent  │ │
          │ │ Service │ │     │ │ Worker  │ │     │ │ Worker  │ │
          │ └─────────┘ │     │ │  × 3    │ │     │ │  × 3    │ │
          │ ┌─────────┐ │     │ └─────────┘ │     │ └─────────┘ │
          │ │   API   │ │     │ ┌─────────┐ │     │ ┌─────────┐ │
          │ │ Gateway │ │     │ │Executor │ │     │ │Executor │ │
          │ └─────────┘ │     │ │ Worker  │ │     │ │ Worker  │ │
          │ ┌─────────┐ │     │ └─────────┘ │     │ └─────────┘ │
          │ │  Redis  │ │     │             │     │             │
          │ │ (Cache) │ │     │             │     │             │
          │ └─────────┘ │     │             │     │             │
          └─────────────┘     └─────────────┘     └─────────────┘
                    │                   │                   │
                    └───────────────────┼───────────────────┘
                                        │
                    ┌───────────────────┼───────────────────┐
                    ▼                   ▼                   ▼
          ┌─────────────┐     ┌─────────────┐     ┌─────────────┐
          │ PostgreSQL  │     │   Milvus    │     │   Neo4j     │
          │ (主) + (从) │     │  (Cluster)  │     │  (Cluster)  │
          └─────────────┘     └─────────────┘     └─────────────┘
                    │                   │                   │
                    └───────────────────┼───────────────────┘
                                        │
                              ┌─────────┴─────────┐
                              │   Kafka Cluster   │
                              │   (3 Brokers)     │
                              └───────────────────┘
"""
add_code_block(deploy_text)

deploy_details = [
    ('高可用', '所有核心服务 3 副本部署，跨可用区分布。Kafka 和数据库均有冗余。'),
    ('弹性伸缩', 'Agent Worker 基于 Kafka 消费延迟自动 HPA 扩缩容。故障高峰时自动扩容，低谷时缩容节省成本。'),
    ('混合云就绪', 'LLM API 调用走公网（Claude API），敏感数据处理的本地模型部署在 VPC 内。所有数据存储在企业私有云或 VPC 内。'),
    ('GitOps', '所有配置通过 Git + ArgoCD 管理。修改配置需要 PR + Review，自动同步到集群。'),
]
for title, desc in deploy_details:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. 迭代路线图
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. 迭代路线图', level=1)

roadmap_headers = ['阶段', '时间', '目标', '核心交付物']
roadmap_rows = [
    ['Phase 0:\n基础设施', 'Month 1-2', '搭建数据管道和基础平台', '• Kafka 告警接入\n• Prometheus/Loki 数据源对接\n• 告警标准化 (CloudEvents)\n• 基础 API Gateway + 认证'],
    ['Phase 1:\n告警分析 MVP', 'Month 2-4', '实现告警自动分类和\n初步分析，人工闭环', '• Alert Classifier Agent\n• 上下文自动增强\n• 简单 Metrics/Log 分析\n• 告警聚合 & 降噪\n• Web Dashboard'],
    ['Phase 2:\n根因分析', 'Month 4-6', '实现假设驱动的多维度\nRCA，覆盖 Top-10 故障类型', '• RCA Synthesizer Agent\n• 假设生成 + 证据收集\n• 变更关联分析\n• 拓扑感知分析\n• 向量库历史案例检索'],
    ['Phase 3:\n半自动修复', 'Month 6-8', 'L1-L2 风险操作自动执行，\nL3 审批后自动执行', '• Auto-Remediation Executor\n• 安全护栏 (OPA)\n• 审批工作流（IM 集成）\n• 修复效果验证\n• 自动回滚机制'],
    ['Phase 4:\n全自动 + 学习', 'Month 8-10', '扩展故障覆盖面，\n实现 L3 部分自动决策', '• 更复杂故障场景覆盖\n• Runbook 自动生成\n• 知识库自进化\n• 复盘报告自动生成\n• 告警预测 (ML)'],
    ['Phase 5:\n平台化', 'Month 10-12', '产品化、多集群、\n多租户、SaaS 化', '• 多集群管理\n• 多租户隔离\n• 自定义 Agent 插件\n• SLA 报表\n• 开放 API'],
]
add_table_with_style(roadmap_headers, roadmap_rows)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. 总结与展望
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. 总结与展望', level=1)

add_para('核心洞察：', bold=True)

insights = [
    'LLM 在运维领域的价值不在于"替代人"，而在于"把人的经验规模化"—— 一个资深 SRE 的经验可以在每次故障中复用，而不是只有在他在线时才能发挥作用。',
    'Multi-Agent 架构的本质优势是"分而治之"：每个 Agent 只关注一个维度（指标/日志/变更/拓扑），组合起来产生 emergent intelligence。这比一个大而全的 Prompt 更可靠、更可调试。',
    '生产环境自动化的关键不是"能做多少"，而是"做错时的后果有多可控"。安全护栏和渐进式自动化是这个系统的生命线。',
    '人机协同的最佳模式是：机器做分析和验证（快+全面），人做决策和担责（审慎+有担当）。系统设计要尊重这个边界。',
]
for s in insights:
    p = doc.add_paragraph(f'• {s}')

doc.add_paragraph()
add_para('未来展望：', bold=True)
future = [
    '预测性运维（Predictive Ops）：从"故障后响应"进化到"故障前预防"，通过 ML 预测即将发生的故障并提前自动修复。',
    '跨系统因果推理：当故障跨越多个系统（网络 + K8s + 数据库 + 应用）时，Agent 能自动构建全局因果图。',
    '自然语言运维：运维人员用自然语言描述问题，Agent 自动完成全流程诊断和修复。"@agent payment 服务很慢，帮我看一下" 即可触发完整分析。',
    '持续学习飞轮：每一次人工干预都成为训练信号 → 模型微调 → Agent 策略优化 → 更少的干预需求 → 最终实现高度自治。',
]
for s in future:
    p = doc.add_paragraph(f'• {s}')

doc.add_paragraph()
doc.add_paragraph()

# 结尾
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— END —')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.bold = True

# ── 保存 ──
output_dir = Path(__file__).resolve().parents[2] / 'docs' / 'design'
output_dir.mkdir(parents=True, exist_ok=True)
output_path = output_dir / '生产环境故障分析自动修复Agent-实现思路.docx'
doc.save(output_path)
print(f'[OK] Document saved: {output_path}')
