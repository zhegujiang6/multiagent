"""Generate Word document: Memory / Planning / Tool Use / Action deep analysis."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from pathlib import Path


def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_code_block(doc, code_text):
    """Add a styled code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Set CJK font too
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2F5496')

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, 'D6E4F0')

    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.4
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ── Title ──
    title = doc.add_heading('多智能体系统核心架构深度解析', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('记忆 · 规划 · 工具使用 · 行动')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    run.bold = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'生成日期：{datetime.date.today().isoformat()}').font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 总体架构概览
    # ════════════════════════════════════════════════
    doc.add_heading('零、四维架构总览', level=1)

    doc.add_paragraph(
        '在 AI Agent 系统中，四个核心能力构成了智能体的闭环运作：\n'
        '• 记忆（Memory）—— 系统"记住"了什么，如何存储和检索信息\n'
        '• 规划（Planning）—— 如何将复杂任务分解为可执行的步骤\n'
        '• 工具使用（Tool Use）—— 系统能调用哪些外部能力来扩展自身\n'
        '• 行动（Action）—— 系统最终产出了什么，改变了什么状态\n\n'
        '本项目在这四个维度上各有精心的工程实现，下面逐一深入分析。'
    )

    overview_rows = [
        ['记忆 (Memory)', '存储层', 'PostgreSQL + Qdrant + Redis + AgentState',
         '会话记忆、知识记忆、画像记忆、工作记忆'],
        ['规划 (Planning)', '编排层', 'LangGraph StateGraph + LLM 路由 + 规则回退',
         '静态骨架 + 动态决策 + 容错回退'],
        ['工具使用 (Tool Use)', '能力层', 'LLM API / PostgreSQL / Qdrant / Redis / Embedding',
         '每个 Agent 硬编码调用特定工具集'],
        ['行动 (Action)', '执行层', '生成回复、创建工单、升级转人工、自进化',
         '状态变更 + 副作用 + 通知推送'],
    ]
    add_styled_table(doc, ['维度', '项目对应层', '核心技术', '说明'], overview_rows)

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 一、记忆 (Memory)
    # ════════════════════════════════════════════════
    doc.add_heading('一、记忆（Memory）', level=1)

    doc.add_paragraph(
        '记忆是 Agent 系统的基础。没有记忆，每次对话都是"失忆"的，系统无法利用历史信息。'
        '本项目实现了多层记忆体系，从毫秒级的缓存到永久存储，覆盖了不同类型的信息。'
    )

    # 1.1 记忆分层架构
    doc.add_heading('1.1 记忆分层架构总览', level=2)

    mem_rows = [
        ['工作记忆\n(Working Memory)', 'AgentState\n(TypedDict)', '单次 Pipeline 执行期间',
         '意图、情绪、画像、工单草稿、路由决策',
         'LangGraph 自动管理；\n每个 Agent 写入，下游 Agent 读取'],
        ['短期记忆\n(Short-term)', 'PostgreSQL\nmessages 表', '最近 10 条消息',
         '对话历史上下文',
         'get_messages(limit=10)；\n角色映射 (customer→user)'],
        ['会话记忆\n(Episodic)', 'PostgreSQL\nconversations 表', '整个会话生命周期',
         '会话状态、情绪趋势 (最近 20 条)、元数据',
         '每轮对话更新 sentiment_trend；\nclosed_at 标记结束'],
        ['语义记忆\n(Semantic)', 'Qdrant 向量库 +\nPostgreSQL knowledge_articles', '持久',
         '50+ FAQ、10 SOP、历史工单解决方案',
         'embed_text() → Qdrant 检索；\nmulti_path_retrieve() 多路融合'],
        ['画像记忆\n(Identity)', 'PostgreSQL\nusers 表', '持久',
         '用户等级 (VIP/Premium/Standard)、标签、历史',
         'ProfileEnricher 按 customer_id 查询'],
        ['审计记忆\n(Audit Trail)', 'PostgreSQL\nagent_runs 表', '持久',
         '每次 Agent 执行的输入/输出/耗时/错误',
         'log_agent_run() 异步记录'],
        ['缓存记忆\n(Cache)', 'Redis', '会话期间',
         '会话状态缓存、Pub/Sub 通道',
         'SLA 升级通知通过 Redis Pub/Sub 推送'],
    ]
    add_styled_table(doc,
        ['记忆层', '存储介质', '生命周期', '存储内容', '读写方式'],
        mem_rows)

    # 1.2 工作记忆
    doc.add_heading('1.2 工作记忆：AgentState — 所有 Agent 的共享"大脑"', level=2)

    doc.add_paragraph(
        'AgentState 是整个系统最核心的数据结构。它是一个 TypedDict，通过 LangGraph 的 '
        'StateGraph 在所有 Agent 节点之间流转。每个 Agent 读取其中的相关字段，执行完毕后'
        '返回部分更新（partial update），LangGraph 自动合并到 State 中传给下一个节点。'
    )

    doc.add_paragraph('AgentState 关键字段与读写关系：', style='List Bullet')

    state_rows = [
        ['messages', 'list[dict]', 'API 层', '所有 Agent',
         '对话历史（最近 10 条），LangGraph add_messages 自动追加'],
        ['conversation_id', 'str', 'API 层', 'FAQ/Ticket/Profile',
         '用于数据库查询和日志关联'],
        ['customer_id', 'str', 'API 层', 'Profile Enricher',
         '用于查询用户画像'],
        ['intent', 'IntentResult', 'Intent Classifier', 'Orchestrator / FAQ / Ticket',
         '意图标签 + 置信度 + 实体列表'],
        ['sentiment', 'SentimentResult', 'Sentiment Analyzer', 'Orchestrator / Ticket / Synthesize',
         '情绪标签 + 触发词 + 趋势评估'],
        ['profile', 'ProfileResult', 'Profile Enricher', 'FAQ / Ticket / Synthesize',
         '用户等级 + 标签 + 服务策略'],
        ['retrieved_knowledge', 'list', 'FAQ Agent', 'Synthesize',
         'RAG 检索到的知识条目'],
        ['response', 'str', '各叶子节点', 'Synthesize → 返回用户',
         '最终回复文本'],
        ['should_escalate', 'bool', 'Sentiment / FAQ / Orchestrator', '路由决策',
         '是否触发转人工'],
        ['ticket_draft', 'TicketDraft', 'Ticket Agent', 'API 层 (创建工单)',
         '工单草稿：标题、分类、优先级、描述'],
        ['agent_decisions', 'list[dict]', '所有 Agent', 'API 层 (审计)',
         '完整的决策链路记录'],
    ]
    add_styled_table(doc,
        ['字段', '类型', '写入者', '读取者', '用途'],
        state_rows)

    doc.add_paragraph(
        '关键设计：AgentState 使用 total=False，意味着所有字段都是可选的。每个 Agent 只填充'
        '自己负责的字段，不干扰其他字段。LangGraph 的 add_messages 注解确保 messages 字段'
        '以追加（而非覆盖）的方式合并。'
    )

    # 1.3 对话记忆
    doc.add_heading('1.3 对话记忆：消息历史的"滑动窗口"', level=2)

    doc.add_paragraph(
        '系统不是无脑地把全部历史塞给 LLM，而是采用"滑动窗口"策略：\n\n'
        '检索策略：\n'
        '• get_messages(conversation_id, limit=10) → 取最近 10 条消息\n'
        '• 按 created_at DESC 查询 → 反转 → 得到时间顺序的最近 10 条\n'
        '• 角色映射：customer → user, agent → assistant（适配 LLM 消息格式）\n\n'
        '为什么是 10 条？\n'
        '• Token 成本控制：10 条消息约 2000-4000 tokens，在 LLM 上下文窗口内\n'
        '• 信息密度：大部分客服对话的最近 5 轮交互足以覆盖当前问题的上下文\n'
        '• 响应速度：更少的上下文 = 更快的 LLM 推理'
    )

    # 1.4 情绪记忆
    doc.add_heading('1.4 情绪记忆：sentiment_trend — 感知用户情绪变化', level=2)

    doc.add_paragraph(
        '系统不只是看"当前"情绪，而是追踪整个会话的情绪变化轨迹：\n\n'
        '存储位置：conversations.sentiment_trend (JSONB 数组)\n'
        '保留策略：最近 20 条 (trend[-20:])\n\n'
        '使用场景：\n'
        '• Sentiment Analyzer 将历史情绪作为上下文传给 LLM，判断趋势\n'
        '• 坐席工作台展示情绪曲线，帮助人工客服了解客户情绪变化\n'
        '• 趋势 declining → 触发预警，提醒坐席注意安抚策略\n\n'
        '趋势评估输出：improving（改善中）/ stable（稳定）/ declining（恶化中）/ first_message（首次）'
    )

    # 1.5 语义记忆
    doc.add_heading('1.5 语义记忆：Qdrant 向量知识库', level=2)

    doc.add_paragraph(
        '这是系统的"长期知识记忆"，存储了所有可被检索和引用的知识。\n\n'
        '知识来源：\n'
        '• 种子数据：50+ FAQ + 10 SOP（seed_data.py）\n'
        '• 自进化：工单解决后自动提取的 Q&A 对\n'
        '• 人工填充：专家填补的知识缺口\n\n'
        '4 个 Collection：\n'
        '• faq_articles — 常见问题\n'
        '• product_docs — 产品文档\n'
        '• sop_documents — 标准操作流程\n'
        '• ticket_resolutions — 历史工单解决方案\n\n'
        '检索方式：\n'
        '① Query Rewriting (LLM 改写口语化问题)\n'
        '② 原始 + 改写双路并行检索\n'
        '③ 按 intent 决定搜索范围 (退款→搜 SOP，技术→搜工单方案)\n'
        '④ 结果融合去重 (按 title 去重，按 score 排序)'
    )

    # 1.6 审计记忆
    doc.add_heading('1.6 审计记忆：agent_runs — 完整执行追溯', level=2)

    doc.add_paragraph(
        '每次 Agent 执行都会记录到 agent_runs 表，包含：\n'
        '• agent_name：哪个 Agent 执行了\n'
        '• input_summary / output_summary：输入输出摘要\n'
        '• latency_ms：执行耗时\n'
        '• tokens_used / model_used：消耗的 Token 和使用的模型\n'
        '• error：如果失败了，错误信息是什么\n\n'
        '这些数据支撑：\n'
        '• 性能监控：哪些 Agent 慢、哪些模型贵\n'
        '• 故障排查：某个会话出了问题，回溯每个 Agent 的输入输出\n'
        '• 成本核算：按 Agent / 模型 / 会话维度统计 Token 消耗'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 二、规划 (Planning)
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('二、规划（Planning）', level=1)

    doc.add_paragraph(
        '规划决定了系统"先做什么、后做什么、遇到什么情况走哪条路"。本项目的规划体系是'
        '"静态骨架 + 动态决策 + 容错回退"的三层混合架构。'
    )

    # 2.1 静态骨架
    doc.add_heading('2.1 第一层：静态骨架 — LangGraph StateGraph', level=2)

    doc.add_paragraph(
        'LangGraph 的 StateGraph 定义了整个处理流程的"骨架"。这个骨架在系统启动时'
        '一次性编译（graph.compile()），之后每次执行都是同一张图的实例化：'
    )

    doc.add_paragraph(
        '静态骨架 = 9 个节点 + 固定的边：\n\n'
        'Node 1: parallel_preprocess    ← 入口点\n'
        '  ↓ (固定边)\n'
        'Node 2: enrich_profile\n'
        '  ↓ (固定边)\n'
        'Node 3: route                  ← 决策点\n'
        '  ↓ (条件边 — 4 条分支)\n'
        '  ├── Node 4a: faq_answer\n'
        '  ├── Node 4b: create_ticket\n'
        '  ├── Node 4c: escalate_to_human\n'
        '  └── Node 4d: direct_response\n'
        '  ↓ (固定边 — 全部汇聚)\n'
        'Node 5: synthesize_response\n'
        '  ↓ (固定边)\n'
        'END'
    )

    doc.add_paragraph(
        '这个骨架的设计原则：\n'
        '• 必须做的（preprocess + enrich）→ 无条件执行\n'
        '• 决策点（route）→ 条件分支\n'
        '• 互斥路径（4 选 1）→ 保证每次只走一条路径\n'
        '• 汇聚点（synthesize）→ 无论哪条路径，最后都要统一包装回复\n\n'
        '编译后的图 = 确定性的执行计划，不会"跑偏"。'
    )

    # 2.2 动态决策
    doc.add_heading('2.2 第二层：动态决策 — LLM 路由 + 规则哨兵', level=2)

    doc.add_paragraph(
        '静态骨架中的 route 节点是整个规划体系的"大脑"，它决定走哪条分支。'
        '路由决策是三段式结构：'
    )

    doc.add_heading('2.2.1 哨兵检查（规则先行，零延迟）', level=3)
    doc.add_paragraph(
        '在调用 LLM 之前，先做零成本的规则检查：\n\n'
        '① 问候哨兵：消息是"你好"/"在吗"等 10 个预设问候词 → 直接走 direct_response\n'
        '   → 无需 LLM，零延迟，零 Token 消耗\n\n'
        '② 升级哨兵：should_escalate 已被设为 True → 直接走 escalate_to_human\n'
        '   → Sentiment Analyzer 已检测到 angry/desperate 时触发'
    )

    doc.add_heading('2.2.2 LLM 路由（智能决策）', level=3)
    doc.add_paragraph(
        '哨兵未命中时，调用 LLM 进行路由决策：\n\n'
        '输入上下文：\n'
        '  {\n'
        '    "intent": {"label": "refund", "confidence": 0.92, ...},\n'
        '    "sentiment": {"label": "dissatisfied", "triggers": [], ...}\n'
        '  }\n\n'
        'LLM 决策规则（写在 System Prompt 中）：\n'
        '• 情绪 angry/desperate → escalate_to_human\n'
        '• 意图 complaint + 法律威胁词 → escalate_to_human\n'
        '• 意图 faq/product_info/account/chitchat → faq_answer\n'
        '• 意图 refund/order_inquiry/tech_support/shipping → create_ticket\n'
        '• 用户说"转人工"/"叫你们经理" → escalate_to_human\n'
        '• 简单问候/感谢 → direct_response\n\n'
        '输出格式（JSON）：{"route": "...", "reason": "..."}'
    )

    doc.add_heading('2.2.3 规则回退（LLM 不可用时的安全保障）', level=3)
    doc.add_paragraph(
        '当 LLM 路由调用失败（JSON 解析异常、网络超时等）时，自动降级为纯规则路由：\n\n'
        '回退规则（优先级从高到低）：\n'
        '1. 情绪 angry/desperate 或有触发词 → escalate_to_human\n'
        '2. 意图 complaint → escalate_to_human\n'
        '3. 意图 faq/product_info/account/chitchat/membership → faq_answer\n'
        '4. 意图 refund/order_inquiry/order_modify/tech_support/shipping/payment → create_ticket\n'
        '5. 以上都不匹配 → faq_answer（默认安全路径）'
    )

    # 2.3 并行规划
    doc.add_heading('2.3 并行规划：asyncio.gather — 时间换效率', level=2)

    doc.add_paragraph(
        '在预处理阶段，Intent Classifier 和 Sentiment Analyzer 通过 asyncio.gather 并行执行：\n\n'
        '串行执行：Intent (3s) → Sentiment (3s) = 6s\n'
        '并行执行：max(Intent 3s, Sentiment 3s) = 3s\n\n'
        '节省了约 50% 的预处理时间。这两个 Agent 之间没有数据依赖（都只读 messages），'
        '天然适合并行化。\n\n'
        'RAG 检索也采用了类似的并行策略：\n'
        '原始 query 检索 (5s) + 改写 query 检索 (5s) = 并行 max(5s, 5s) = 5s'
    )

    # 2.4 超时规划
    doc.add_heading('2.4 超时保护：双层时间预算', level=2)

    timeout_rows = [
        ['单次 LLM 调用', 'BaseAgent._call_llm()', '15s',
         '防止某个 LLM 调用卡死整个流程'],
        ['RAG 检索', 'multi_path_retrieve()', '8s × 2',
         'Qdrant 不可用时快速降级，不阻塞'],
        ['Orchestrator Pipeline', 'orchestrator.run()', '25s',
         '整体流程超时 → 自动转人工'],
        ['Streaming 外层', 'process_message_streaming()', '30s',
         '安全网：防止 Orchestrator 自身超时失效'],
    ]
    add_styled_table(doc,
        ['超时层级', '位置', '阈值', '超时后行为'],
        timeout_rows)

    doc.add_paragraph(
        '设计理念：每一层都有自己的时间预算，上层比下层更宽松（15s < 25s < 30s），'
        '形成"漏斗式"超时保护。'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 三、工具使用 (Tool Use)
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('三、工具使用（Tool Use）', level=1)

    doc.add_paragraph(
        '工具使用扩展了 LLM 的能力边界——LLM 本身只能生成文本，但通过调用外部工具，'
        '系统可以查询数据库、搜索知识库、创建工单、发送通知等。\n\n'
        '本项目的工具使用模式是"硬编码调用"：每个 Agent 的 execute() 方法中硬编码了'
        '它需要调用的工具，而非让 LLM 动态选择工具。这种设计适合当前阶段——Agent 的职责'
        '明确、工具集固定、调用顺序确定。'
    )

    # 3.1 工具矩阵
    doc.add_heading('3.1 工具矩阵：每个 Agent 使用了哪些工具', level=2)

    tool_rows = [
        ['Intent\nClassifier', 'LLM API\n(FAST_MODEL)',
         '对用户消息进行意图分类和实体提取',
         'OpenAI-compatible API\n(qwen-max / haiku)',
         '15s 超时，双格式兼容'],
        ['Sentiment\nAnalyzer', 'LLM API\n(FAST_MODEL)',
         '分析用户情绪和检测触发词',
         '同上',
         '传入历史情绪作为上下文'],
        ['Profile\nEnricher', 'PostgreSQL\n(SQLAlchemy)',
         '查询用户画像（等级/标签/历史）',
         'async_session_factory',
         '失败 → 降级为默认画像'],
        ['FAQ Agent', '① Qdrant (向量检索)\n② LLM API (答案生成)\n③ Query Rewriting (LLM)\n④ PostgreSQL (gap/usage)',
         'RAG 检索→答案生成→知识缺口检测→使用追踪',
         'Qdrant Client +\nOpenAI API +\nSQLAlchemy',
         '每步都有 try/except 和降级'],
        ['Ticket\nAgent', 'LLM API\n(FAST_MODEL)',
         '从对话提取工单信息并计算优先级',
         'OpenAI-compatible API',
         '失败 → 回退为默认工单信息'],
        ['Orchestrator', 'LLM API\n(FAST_MODEL)',
         '路由决策',
         'OpenAI-compatible API',
         '失败 → 规则回退路由'],
        ['SLA Service', '① PostgreSQL\n② Redis Pub/Sub',
         '轮询 SLA 截止时间 → 升级通知',
         'SQLAlchemy +\naioredis',
         '后台 30s 循环'],
        ['Knowledge\nExtraction', 'LLM API\n(DEFAULT_MODEL)',
         '从已解决对话中提取 Q&A 对',
         'OpenAI-compatible API',
         '低 temperature (0.3) 保证稳定性'],
        ['Embedding', '① sentence-transformers (本地)\n② OpenAI Embedding API (远程)',
         '文本向量化',
         '本地模型或远程 API',
         '自动检测配置选择本地/远程'],
    ]
    add_styled_table(doc,
        ['Agent/服务', '使用的工具', '工具用途', '工具实现', '容错机制'],
        tool_rows)

    # 3.2 核心工具详解
    doc.add_heading('3.2 核心工具详解', level=2)

    doc.add_heading('3.2.1 LLM API 工具 — BaseAgent._call_llm()', level=3)
    doc.add_paragraph(
        '这是所有 Agent 共用的核心工具，封装了与 LLM 的交互。不是简单的 HTTP 调用，'
        '而是一个带有完整保护机制的工具：\n\n'
        '① 超时保护：asyncio.wait_for(timeout=15s)\n'
        '② 双格式兼容：同时支持标准 OpenAI 格式 (choices[0].message.content) '
        '和阿里 MaaS 格式 (resp.text)\n'
        '③ 结构化输出：System Prompt 要求 LLM 输出严格 JSON，代码做 JSON.parse\n'
        '④ 日志记录：每次调用记录耗时和内容长度\n'
        '⑤ 异常透传：不吞异常，让上层 Agent 自己决定降级策略'
    )

    doc.add_heading('3.2.2 RAG 检索工具 — multi_path_retrieve()', level=3)
    doc.add_paragraph(
        '这是 FAQ Agent 专用的复合工具，整合了多个子工具：\n\n'
        'Step 1 — Query Rewriting（LLM 工具）：\n'
        '  用户输入："你们那个退钱怎么弄"\n'
        '  LLM 改写："退款流程"\n\n'
        'Step 2 — 双路并行检索（Qdrant 工具）：\n'
        '  Path A: embed("你们那个退钱怎么弄") → Qdrant.search()\n'
        '  Path B: embed("退款流程") → Qdrant.search()\n'
        '  并行执行，各 8s 超时\n\n'
        'Step 3 — 结果融合：\n'
        '  合并 A + B → 按 title 去重 → 按 score 降序 → 取 top_k\n\n'
        'Step 4 — PII 脱敏（安全工具）：\n'
        '  mask_pii(content) → 隐藏手机号/身份证/银行卡/邮箱'
    )

    doc.add_heading('3.2.3 数据库工具 — Profile Enricher 的查询模式', level=3)
    doc.add_paragraph(
        'Profile Enricher 使用 SQLAlchemy 异步查询 PostgreSQL：\n\n'
        '查询逻辑：\n'
        '  SELECT * FROM users\n'
        '  WHERE id = :customer_id OR external_id = :customer_id\n\n'
        '支持两种 ID 匹配：内部 UUID 和外部系统 ID（如 CRM 系统的客户编号）。\n'
        '查询失败 → 降级为标准访客画像，不阻塞流程。'
    )

    doc.add_heading('3.2.4 Embedding 工具 — 双模嵌入', level=3)
    doc.add_paragraph(
        '系统支持两种 Embedding 模式，通过环境变量自动选择：\n\n'
        '模式 A — 本地模型（默认）：\n'
        '  sentence-transformers / paraphrase-multilingual-MiniLM-L12-v2\n'
        '  维度：384，免费，离线可用\n\n'
        '模式 B — 远程 API：\n'
        '  配置 EMBEDDING_API_KEY + EMBEDDING_API_BASE 后自动切换\n'
        '  支持 OpenAI / 阿里等兼容 API\n'
        '  维度可配置（如 text-embedding-v4 的 1024 维）\n\n'
        '选择逻辑：_should_use_api() → 检查是否配置了 API Key 和 Base URL'
    )

    doc.add_heading('3.2.5 Redis 工具 — 发布订阅 + 缓存', level=3)
    doc.add_paragraph(
        'Redis 承担两个角色：\n\n'
        '① 发布订阅（Pub/Sub）：\n'
        '  SLA 升级通知 → publish("sla:escalations", msg)\n'
        '  前端坐席工作台订阅该频道，实时收到升级提醒\n\n'
        '② 会话缓存：\n'
        '  活跃会话状态缓存在 Redis 中\n'
        '  WebSocket 连接状态管理'
    )

    doc.add_heading('3.3 工具使用的设计模式', level=2)
    doc.add_paragraph(
        '本项目采用的工具使用模式可以总结为"硬编码工具调用"（Hardcoded Tool Calling），'
        '与当前流行的"LLM 动态函数调用"（Function Calling）形成对比：\n\n'
        '本项目的模式：\n'
        '  Agent.execute(state) {\n'
        '    result = await tool.call(state)  // 硬编码调用\n'
        '    return {"field": result}\n'
        '  }\n\n'
        'Function Calling 模式：\n'
        '  LLM 决定: "我需要调用 search_knowledge_base 工具"\n'
        '  系统执行: search_knowledge_base(query)\n'
        '  LLM 处理: 根据返回结果继续推理\n\n'
        '本项目的选择是合理的——当前 6 个 Agent 的职责和工具集都是明确的，不需要 LLM '
        '动态选择工具。如果未来 Agent 数量增加或工具集扩大，可以考虑引入 Function Calling。'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 四、行动 (Action)
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('四、行动（Action）', level=1)

    doc.add_paragraph(
        '行动是系统的"输出端"——经过感知、记忆检索和规划后，系统最终做了什么。'
        '本项目的行动可以分为五大类：回复行动、工单行动、升级行动、知识进化行动和监控行动。'
    )

    # 4.1 回复行动
    doc.add_heading('4.1 回复行动：生成用户可读的响应', level=2)

    doc.add_paragraph(
        '这是最直接、最频繁的行动。系统根据不同路径生成不同风格的回复：'
    )

    reply_rows = [
        ['FAQ 回答', 'FAQ Agent', 'RAG 检索结果 + LLM 生成',
         '"根据知识库中的信息：退换货需要在签收后 7 天内申请..."\n【参考来源：知识1 ｜ 相关度：0.92】'],
        ['工单回复', 'Orchestrator\n_create_ticket_node', '工单草稿 + SLA 信息拼装',
         '"我已收到您的问题，为您创建了工单。\n📋 优先级 P2，预计 8 小时内处理..."'],
        ['转人工回复', 'Orchestrator\n_escalate_node', '固定模板',
         '"已升级为人工服务，客服人员正在赶来，预计等待不超过 30 秒..."'],
        ['问候回复', 'Orchestrator\n_direct_response_node', '固定模板',
         '"您好！我是智能客服助手，请问有什么可以帮助您的吗？"'],
        ['知识缺口回复', 'FAQ Agent', '检索得分 < 0.5 触发',
         '"抱歉，我暂时没有找到足够相关的信息。请告诉我更多细节..."'],
        ['超时降级回复', 'Orchestrator.run()', 'Pipeline 超时 25s 触发',
         '"抱歉，系统处理您的请求时超时了，正在为您转接人工客服..."'],
    ]
    add_styled_table(doc,
        ['回复类型', '触发位置', '生成方式', '示例'],
        reply_rows)

    doc.add_heading('4.1.1 回复增强：共情包装', level=3)
    doc.add_paragraph(
        'SynthesizeResponse 节点对所有回复做最终包装：\n\n'
        '① 共情前缀（基于情绪）：\n'
        '  angry    → "非常理解您的心情，给您带来不好的体验我们深表歉意。"\n'
        '  dissatisfied → "很抱歉给您带来不便，让我来帮您处理。"\n'
        '  desperate → "您的问题我们已经高度重视，我将立即为您升级处理。"\n\n'
        '② VIP 后缀（基于画像）：\n'
        '  tier=vip → "💎 作为我们的 VIP 会员，您的问题享有最高优先级服务。"\n\n'
        '③ 去重保护：如果原始回复已包含共情语句，不重复添加'
    )

    # 4.2 工单行动
    doc.add_heading('4.2 工单行动：创建和管理工单', level=2)

    doc.add_paragraph(
        '工单行动是最复杂的行动类型，涉及状态变更、SLA 计算和跨部门流转：'
    )

    doc.add_heading('4.2.1 工单创建', level=3)
    doc.add_paragraph(
        'Ticket Agent 提取工单信息后，由 Orchestrator 的 _create_ticket_node 组装回复，'
        '但实际的数据库写入由 API 层的 process_message() 完成。\n\n'
        '创建工单时执行的行动：\n'
        '① 生成 display_id：TK-20240701-0001 格式\n'
        '② 计算 SLA 截止时间：response_deadline + resolve_deadline\n'
        '③ 写入 tickets 表\n'
        '④ 记录 ticket_events（from_status=None → to_status="new"）\n'
        '⑤ WebSocket 推送 {"type": "ticket_created", "draft": {...}}'
    )

    doc.add_heading('4.2.2 工单状态转换', level=3)
    doc.add_paragraph(
        '严格的状态机校验，非法转换直接拒绝：\n\n'
        'VALID_TRANSITIONS = {\n'
        '  "new":       → ["assigned"],\n'
        '  "assigned":  → ["in_progress"],\n'
        '  "in_progress": → ["pending", "waiting", "resolved"],\n'
        '  "pending":   → ["in_progress", "resolved"],\n'
        '  "waiting":   → ["in_progress", "resolved"],\n'
        '  "resolved":  → ["closed", "reopened"],\n'
        '  "closed":    → ["reopened"],\n'
        '  "reopened":  → ["in_progress"],\n'
        '}\n\n'
        '每次状态变更都记录到 ticket_events 表，形成完整的审计追踪。'
    )

    # 4.3 升级行动
    doc.add_heading('4.3 升级行动：多层次的自动干预', level=2)

    esc_rows = [
        ['情绪升级', 'Sentiment Analyzer\n检测到 angry/desperate',
         '设置 should_escalate=True\n路由直接走 escalate_to_human',
         '即时（预处理阶段）'],
        ['路由升级', 'Orchestrator LLM\n路由决策',
         'complaint + 法律威胁词 → escalate_to_human\n用户说"转人工" → escalate_to_human',
         '路由阶段'],
        ['SLA 预警', 'sla_service 后台轮询',
         '50% 时间 → 标记 sla_warning_sent\n80% 时间 → 升级优先级 + 重新计算 SLA\n100% 时间 → 紧急升级 P0 + Redis 通知',
         '后台 30s 轮询'],
        ['超时升级', 'orchestrator.run()\n超时 25s',
         '自动转人工 + 返回超时提示',
         '异常处理阶段'],
    ]
    add_styled_table(doc,
        ['升级类型', '触发条件', '执行动作', '时机'],
        esc_rows)

    # 4.4 知识进化行动
    doc.add_heading('4.4 知识进化行动：自主学习和知识增长', level=2)

    doc.add_paragraph(
        '这是系统最具特色的行动类型——不需要人工干预，系统自动完成知识的发现、提取、审核和发布。'
    )

    evo_rows = [
        ['知识缺口\n检测', 'FAQ Agent 检索\nbest_score < 0.5',
         '后台异步调用 create_gap_record()',
         '"用户问了什么但知识库里没有" → 记录到 knowledge_articles (status=gap)'],
        ['知识使用\n追踪', 'FAQ Agent 回答成功后',
         '后台异步调用 track_knowledge_usage()',
         '更新 usage_count++ 和 effectiveness_score (滚动平均)'],
        ['知识自动\n提取', '工单 resolved 时',
         '后台异步调用 extract_knowledge_from_conversation()',
         'LLM 从对话中提取 Q&A 对'],
        ['自动审核\n分级发布', '知识提取完成后',
         'confidence ≥ 0.8 → auto_approve + Qdrant\n0.6-0.8 → draft 待人工审核\n< 0.6 → 删除（低价值）',
         '高置信度全自动，低置信度人工兜底'],
        ['人工填坑', '专家通过 Knowledge UI',
         'fill_gap() → 自动审批 + 发布到 Qdrant',
         '填补 AI 发现的知识空白'],
    ]
    add_styled_table(doc,
        ['行动名称', '触发时机', '具体动作', '效果'],
        evo_rows)

    doc.add_heading('4.4.1 知识进化闭环全景', level=3)
    doc.add_paragraph(
        '┌──────────────────────────────────────────────────────────────┐\n'
        '│                                                              │\n'
        '│   用户提问 ──→ RAG 检索 ──→ 检索得分 < 0.5？                │\n'
        '│                              │                               │\n'
        '│                     YES ↓           NO ↓                     │\n'
        '│                  记录 Gap         生成回答                    │\n'
        '│                     │               │                        │\n'
        '│              人工专家审核         追踪 usage                  │\n'
        '│                     │               │                        │\n'
        '│                     └───────┬───────┘                        │\n'
        '│                             ↓                                │\n'
        '│                     发布到 Qdrant                            │\n'
        '│                                                              │\n'
        '│   ─────────────────── 另一条进化路径 ───────────────────     │\n'
        '│                                                              │\n'
        '│   工单 resolved ──→ LLM 提取 Q&A ──→ confidence ≥ 0.8？      │\n'
        '│                                       │                      │\n'
        '│                              YES ↓          NO ↓             │\n'
        '│                          自动发布      存为 Draft            │\n'
        '│                                         (人工审核)           │\n'
        '│                                                              │\n'
        '└──────────────────────────────────────────────────────────────┘'
    )

    # 4.5 监控行动
    doc.add_heading('4.5 监控行动：后台持续运行的守护进程', level=2)

    doc.add_paragraph(
        '系统启动时（FastAPI lifespan），自动启动两个后台任务：\n\n'
        '① SLA 轮询（_sla_polling）：\n'
        '  每 30 秒查询所有活跃工单 → 检查 SLA 消耗比例 → 预警/升级/紧急通知\n'
        '  通过 Redis Pub/Sub 实时推送给前端坐席工作台\n\n'
        '② WebSocket 心跳：\n'
        '  每 30 秒 ping/pong 保持连接\n'
        '  连接断开 → 自动清理会话状态'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 五、四维联动
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('五、四维联动：一次完整请求的全景追踪', level=1)

    doc.add_paragraph(
        '下面以"VIP 用户退货"场景为例，展示记忆、规划、工具使用和行动是如何协同工作的：'
    )

    doc.add_paragraph(
        '用户消息："我上个月买的这个破东西，用了没几天就坏了，我要退钱！"\n\n'
        '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n'
        '【记忆 — 读取】\n'
        '  1. 从 PostgreSQL 读取最近 10 条消息 → 构建 AgentState.messages\n'
        '  2. 从 PostgreSQL 查询用户画像 → tier=vip, tags=["高价值"]\n'
        '  3. 历史情绪趋势 → ["dissatisfied", "neutral"] → 趋势 declining\n\n'
        '【规划 — 并行预处理】\n'
        '  4. Intent Classifier → "refund", confidence=0.85, entities=[{"商品",""}]\n'
        '  5. Sentiment Analyzer → "angry", triggers=["退钱"+"马上"暗示]\n'
        '  6. Profile Enricher → tier=vip, suggested_tone="热情主动"\n\n'
        '【规划 — 路由决策】\n'
        '  7. 哨兵检查：should_escalate? → 否（angry 但无法律威胁词）\n'
        '  8. LLM 路由：intent=refund + sentiment=angry + tier=vip\n'
        '     → 决策：create_ticket（退款需要工单，但升级优先级）\n\n'
        '【工具使用 — 创建工单】\n'
        '  9. Ticket Agent 调用 LLM → 提取工单信息\n'
        '     标题："VIP用户申请商品退货退款"\n'
        '     分类：refund\n'
        '     优先级：P2 → angry 升级为 P1 → VIP 再升级为 P0\n'
        '     SLA：响应 15 分钟 / 解决 1 小时\n\n'
        '【行动 — 生成回复】\n'
        '  10. Synthesize 添加共情前缀："非常理解您的心情..."\n'
        '  11. Synthesize 添加 VIP 后缀："💎 作为 VIP 会员，您的问题享有最高优先级..."\n'
        '  12. WebSocket 推送 ticket_created 事件\n\n'
        '【记忆 — 写入】\n'
        '  13. 保存 Agent 回复到 messages 表\n'
        '  14. 记录 agent_runs (6 个 Agent 的执行日志)\n'
        '  15. 更新 conversation.sentiment_trend 追加 "angry"\n'
        '  16. 更新 conversation.updated_at\n\n'
        '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
        '总耗时：约 8-10 秒'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 六、设计评价
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('六、设计评价与改进方向', level=1)

    doc.add_heading('6.1 当前架构的优势', level=2)

    advantages = [
        ('记忆分层清晰',
         '从 AgentState（毫秒级）→ Redis（秒级）→ PostgreSQL（持久）→ Qdrant（语义），'
         '每层有明确的职责和生命周期，不会互相污染。'),
        ('规划可预测',
         'LangGraph 的编译时图 + 条件边保证了执行路径的可预测性和可调试性。'
         '任何会话都可以通过 agent_decisions 回溯完整的决策链路。'),
        ('容错覆盖全面',
         'LLM 超时、DB 连接失败、Qdrant 不可用、Pipeline 整体超时——每个故障点都有降级策略，'
         '不会让用户"干等"或"报错"。'),
        ('工具职责单一',
         '每个 Agent 只调用自己需要的工具，工具和 Agent 的对应关系清晰，易于测试和维护。'),
        ('行动可审计',
         '所有状态变更都记录到数据库（messages、tickets、ticket_events、agent_runs），'
         '任何问题都可以追溯。'),
    ]
    for title, desc in advantages:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}：')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('6.2 可改进的方向', level=2)

    improvements = [
        ('引入动态工具选择（Function Calling）',
         '当前工具调用是硬编码的。如果未来 Agent 数量增加，可以考虑让 LLM 动态选择工具。'
         '但这会增加不确定性，需要权衡。'),
        ('记忆窗口自适应',
         '当前固定取最近 10 条消息。可以根据对话复杂度动态调整——简单问候取 2 条，'
         '复杂投诉取 20 条。'),
        ('增加反思（Reflection）环节',
         '当前是"一次性"生成回复。可以增加自我检查环节：生成回复后，再让 LLM 评估'
         '回复是否准确、完整、有共情，不通过则重新生成。'),
        ('长期用户画像记忆',
         '当前用户画像只从数据库读取静态字段。可以增加从历史对话中自动提取的画像信息——'
         '例如"该用户对物流时效特别敏感"、"偏好电话沟通"。'),
        ('跨会话记忆',
         '当前记忆以单次会话为边界。如果用户隔天再来，系统"不认识"他。可以通过'
         '全局用户向量记忆解决这个问题。'),
    ]
    for title, desc in improvements:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title}：')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 文档结束 —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    return doc


if __name__ == '__main__':
    doc = build_document()
    output_dir = Path(__file__).resolve().parents[2] / 'docs' / 'design'
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / '多智能体系统-记忆规划工具行动-深度分析.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')
